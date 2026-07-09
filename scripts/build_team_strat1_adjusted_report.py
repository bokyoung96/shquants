from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

import matplotlib.pyplot as plt
from matplotlib.patches import Rectangle
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import landscape, letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from PIL import Image as PILImage


ROOT = Path(__file__).resolve().parents[1]
PARQUET = ROOT / "parquet"
OUT_DIR = ROOT / "results" / "team_strat1"
FIG_DIR = OUT_DIR / "adjusted_figures"
REPORT_STEM = "report_adjusted"
NORMAL_TRADING_LABEL = "정상"
ANALYSIS_START_DATE = pd.Timestamp("2016-01-01")
ANALYSIS_END_DATE = pd.Timestamp("2026-05-27")

HORIZONS = tuple(range(1, 6))
PERIOD_SPECS = (
    ("2016~현재", "2016-01-01"),
)

BLUE = RGBColor(31, 77, 120)
DARK = RGBColor(32, 43, 54)
MUTED = RGBColor(98, 110, 125)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "B8C2CC"

NAVY = "#102A43"
INK = "#1F2937"
SLATE = "#64748B"
LINE = "#D8E0EA"
PANEL = "#F6F8FB"
TEAL = "#0E7490"
CYAN = "#38BDF8"
CORAL = "#EF6F61"
GOLD = "#D6A84F"
GREEN = "#14B8A6"
GRAY_BAR = "#CBD5E1"


@dataclass(frozen=True)
class ReportData:
    events: pd.DataFrame
    summary: pd.DataFrame
    long_summary: pd.DataFrame
    counts: pd.DataFrame
    path_excluded: pd.DataFrame
    latest_price_date: pd.Timestamp
    latest_signal_date: pd.Timestamp
    figures: dict[str, Path]


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    data = build_report_data()
    write_workbook(data)
    docx_path = OUT_DIR / f"{REPORT_STEM}.docx"
    build_docx(data, docx_path)
    pdf_path = OUT_DIR / f"{REPORT_STEM}.pdf"
    build_pdf(data, pdf_path)
    print(f"wrote {docx_path}")
    print(f"wrote {pdf_path}")
    print(f"wrote {OUT_DIR / f'{REPORT_STEM}.xlsx'}")


def build_report_data() -> ReportData:
    open_ = read_price("qw_adj_o")
    high = read_price("qw_adj_h").reindex_like(open_)
    low = read_price("qw_adj_l").reindex_like(open_)
    close = read_price("qw_adj_c").reindex_like(open_)
    volume = read_price("qw_v").reindex_like(open_)
    k200 = read_frame("qw_k200_yn").reindex(index=open_.index, columns=open_.columns).fillna(0).astype(int).eq(1)
    tradable_raw = read_frame("qw_trs_ban").reindex(index=open_.index, columns=open_.columns)
    normal_value = _normal_trading_value(tradable_raw)
    tradable = tradable_raw.eq(normal_value) | tradable_raw.isna()

    body_top = pd.DataFrame(np.maximum(open_.to_numpy(), close.to_numpy()), index=open_.index, columns=open_.columns)
    body_bottom = pd.DataFrame(np.minimum(open_.to_numpy(), close.to_numpy()), index=open_.index, columns=open_.columns)

    range_ok = high.sub(low).divide(close.shift(1)).gt(0.10)
    upper_wick_ok = high.sub(body_top).gt(body_top.sub(body_bottom))
    runup_ok = high.divide(close.shift(63)).gt(1.5)
    volume_rank = volume.rolling(252, min_periods=252).rank(pct=True)
    volume_ok = volume_rank.gt(0.90)
    raw_signal = range_ok & upper_wick_ok & runup_ok & volume_ok
    gated_signal = raw_signal & k200 & tradable

    latest_signal_date = ANALYSIS_END_DATE
    events = deduplicate_events(gated_signal.loc[:latest_signal_date], open_, high, close)
    events = enrich_event_names(events)
    events = events[events["signal_date"].ge(ANALYSIS_START_DATE)].reset_index(drop=True)
    events, path_excluded = apply_path_normal_filter(events, open_, high, low, close, tradable_raw, normal_value)
    summary = summarize_periods(events, latest_signal_date, latest_signal_date)
    long_summary = make_long_summary(summary)
    counts = make_counts(raw_signal, gated_signal, events, latest_signal_date, ANALYSIS_START_DATE, path_excluded)
    figures = make_figures(summary, long_summary, counts, events, open_, high, low, close)
    return ReportData(events, summary, long_summary, counts, path_excluded, latest_signal_date, latest_signal_date, figures)


def read_frame(stem: str) -> pd.DataFrame:
    return pd.read_parquet(PARQUET / f"{stem}.parquet")


def read_price(stem: str) -> pd.DataFrame:
    return read_frame(stem).astype(float)


def _normal_trading_value(frame: pd.DataFrame) -> object:
    values = list(pd.unique(frame.stack().dropna()))
    for value in values:
        if str(value) == NORMAL_TRADING_LABEL:
            return value
    if values:
        return values[0]
    return NORMAL_TRADING_LABEL


def apply_path_normal_filter(
    events: pd.DataFrame,
    open_: pd.DataFrame,
    high: pd.DataFrame,
    low: pd.DataFrame,
    close: pd.DataFrame,
    tradable_raw: pd.DataFrame,
    normal_value: object,
) -> tuple[pd.DataFrame, pd.DataFrame]:
    rows: list[dict[str, object]] = []
    keep: list[bool] = []
    for _, event in events.iterrows():
        signal_date = pd.Timestamp(event["signal_date"])
        symbol = str(event["symbol"])
        pos = open_.index.get_loc(signal_date)
        reasons: list[str] = []
        abnormal_dates: list[str] = []
        for check_pos in range(pos + 1, pos + 6):
            check_date = open_.index[check_pos]
            status = tradable_raw.at[check_date, symbol] if check_date in tradable_raw.index and symbol in tradable_raw.columns else np.nan
            is_halt = pd.notna(status) and status != normal_value
            o = open_.at[check_date, symbol]
            h = high.at[check_date, symbol]
            low_price = low.at[check_date, symbol]
            c = close.at[check_date, symbol]
            all_equal = (
                pd.notna(o)
                and pd.notna(h)
                and pd.notna(low_price)
                and pd.notna(c)
                and float(o) == float(h) == float(low_price) == float(c)
            )
            if is_halt or all_equal:
                reason_bits = []
                if is_halt:
                    reason_bits.append(f"status={status}")
                if all_equal:
                    reason_bits.append("open=high=low=close")
                abnormal_dates.append(f"T+{check_pos - pos} {check_date:%Y-%m-%d} ({', '.join(reason_bits)})")
        if abnormal_dates:
            reasons.append("; ".join(abnormal_dates))
        keep.append(not reasons)
        if reasons:
            rows.append(
                {
                    "signal_date": signal_date,
                    "symbol": symbol,
                    "name": event.get("name", ""),
                    "entry_date": event.get("entry_date"),
                    "path_normal_exclusion_reason": " | ".join(reasons),
                }
            )
    excluded = pd.DataFrame(rows)
    filtered = events.loc[keep].reset_index(drop=True).copy()
    filtered["path_normal"] = True
    return filtered, excluded


def deduplicate_events(mask: pd.DataFrame, open_: pd.DataFrame, high: pd.DataFrame, close: pd.DataFrame) -> pd.DataFrame:
    suppress_until: dict[str, int] = {}
    events: list[dict[str, object]] = []
    index = open_.index
    for local_pos, signal_date in enumerate(mask.index):
        orig_pos = index.get_loc(signal_date)
        if orig_pos + 5 >= len(index):
            continue
        symbols = list(mask.columns[mask.loc[signal_date].fillna(False).to_numpy(dtype=bool)])
        for symbol in symbols:
            if suppress_until.get(symbol, -1) >= local_pos:
                continue
            entry_date = index[orig_pos + 1]
            entry_open = open_.at[entry_date, symbol]
            signal_high = high.at[signal_date, symbol]
            if pd.isna(entry_open) or pd.isna(signal_high):
                continue
            row: dict[str, object] = {
                "signal_date": signal_date,
                "symbol": symbol,
                "entry_date": entry_date,
                "entry_open": float(entry_open),
                "signal_high": float(signal_high),
                "gap_over_signal_high": bool(float(entry_open) > float(signal_high)),
            }
            row["entry_status"] = "not_entered" if row["gap_over_signal_high"] else "entered"
            first_touch_date = pd.NaT
            first_touch_price = np.nan
            for horizon in HORIZONS:
                exit_date = index[orig_pos + horizon]
                exit_close = close.at[exit_date, symbol]
                row[f"exit_date_T{horizon}"] = exit_date
                row[f"exit_close_T{horizon}"] = float(exit_close) if pd.notna(exit_close) else np.nan
                row[f"T+{horizon}"] = (
                    (float(entry_open) - float(exit_close)) / float(entry_open)
                    if pd.notna(exit_close) and float(entry_open) != 0.0
                    else np.nan
                )
                if row["gap_over_signal_high"]:
                    row[f"prior_high_touch_T{horizon}"] = False
                    row[f"adjusted_exit_date_T{horizon}"] = entry_date
                    row[f"adjusted_exit_price_T{horizon}"] = float(entry_open)
                    row[f"adjusted_exit_reason_T{horizon}"] = "not_entered"
                    row[f"adjusted_T+{horizon}"] = 0.0
                    continue

                touch_date = pd.NaT
                touch_price = np.nan
                fixed_exit_high = float(signal_high)
                for check_pos in range(orig_pos + 1, orig_pos + horizon + 1):
                    check_date = index[check_pos]
                    check_open = open_.at[check_date, symbol]
                    check_high = high.at[check_date, symbol]
                    if pd.notna(check_high) and float(check_high) >= fixed_exit_high:
                        touch_date = check_date
                        if check_pos > orig_pos + 1 and pd.notna(check_open) and float(check_open) > fixed_exit_high:
                            touch_price = float(check_open)
                            adjusted_exit_reason = "gap_open_stop"
                        else:
                            touch_price = fixed_exit_high
                            adjusted_exit_reason = "prior_high_touch"
                        break

                touched = pd.notna(touch_date)
                row[f"prior_high_touch_T{horizon}"] = bool(touched)
                if touched:
                    adjusted_exit_date = touch_date
                    adjusted_exit_price = touch_price
                    if pd.isna(first_touch_date):
                        first_touch_date = touch_date
                        first_touch_price = touch_price
                else:
                    adjusted_exit_date = exit_date
                    adjusted_exit_price = float(exit_close) if pd.notna(exit_close) else np.nan
                    adjusted_exit_reason = "horizon_close"
                row[f"adjusted_exit_date_T{horizon}"] = adjusted_exit_date
                row[f"adjusted_exit_price_T{horizon}"] = adjusted_exit_price
                row[f"adjusted_exit_reason_T{horizon}"] = adjusted_exit_reason
                row[f"adjusted_T+{horizon}"] = (
                    (float(entry_open) - float(adjusted_exit_price)) / float(entry_open)
                    if pd.notna(adjusted_exit_price) and float(entry_open) != 0.0
                    else np.nan
                )
            row["first_prior_high_touch_date"] = first_touch_date
            row["first_prior_high_touch_price"] = first_touch_price
            events.append(row)
            suppress_until[symbol] = local_pos + 5
    return pd.DataFrame(events)


def enrich_event_names(events: pd.DataFrame) -> pd.DataFrame:
    previous = OUT_DIR / "team_strat1_report.xlsx"
    enriched = events.copy()
    enriched["name"] = ""
    enriched["sector"] = ""
    if previous.exists():
        try:
            names = pd.read_excel(previous, sheet_name="EVENTS_358", usecols=["signal_date", "symbol", "name", "sector"])
            names["signal_date"] = pd.to_datetime(names["signal_date"])
            enriched["signal_date"] = pd.to_datetime(enriched["signal_date"])
            enriched = enriched.merge(names, on=["signal_date", "symbol"], how="left", suffixes=("", "_known"))
            enriched["name"] = enriched["name_known"].fillna(enriched["name"]).fillna("")
            enriched["sector"] = enriched["sector_known"].fillna(enriched["sector"]).fillna("")
            enriched = enriched.drop(columns=[c for c in ("name_known", "sector_known") if c in enriched.columns])
        except Exception:
            pass
    name_map_path = PARQUET / "map__ticker_name_gics_sector_map.parquet"
    if enriched["name"].astype(str).eq("").all() and name_map_path.exists():
        name_map = pd.read_parquet(name_map_path).rename(
            columns={"TICKER": "symbol", "NAME": "name_known", "GICS_SECTOR_NAME": "sector_known"}
        )
        enriched = enriched.merge(name_map[["symbol", "name_known", "sector_known"]], on="symbol", how="left")
        enriched["name"] = enriched["name_known"].fillna(enriched["name"]).fillna("")
        enriched["sector"] = enriched["sector_known"].fillna(enriched["sector"]).fillna("")
        enriched = enriched.drop(columns=["name_known", "sector_known"])
    return enriched


def summarize_periods(events: pd.DataFrame, latest_signal_date: pd.Timestamp, latest_price_date: pd.Timestamp) -> pd.DataFrame:
    rows = []
    for period, start in PERIOD_SPECS:
        subset = events if start is None else events[events["signal_date"].ge(pd.Timestamp(start))]
        period_start = pd.Timestamp(start) if start is not None else subset["signal_date"].min()
        row: dict[str, object] = {
            "period": period,
            "start": period_start,
            "end": latest_signal_date,
            "events": int(len(subset)),
            "signal_days": int(subset["signal_date"].nunique()),
            "symbols": int(subset["symbol"].nunique()),
            "veto_events": int(subset["gap_over_signal_high"].sum()),
            "not_entered_events": int(subset["gap_over_signal_high"].sum()),
            "prior_high_touch_events": int(subset["first_prior_high_touch_date"].notna().sum()),
            "veto_rate": float(subset["gap_over_signal_high"].mean()),
            "latest_signal_date": latest_signal_date,
            "latest_price_date": latest_price_date,
        }
        active = subset[~subset["gap_over_signal_high"]]
        for horizon in HORIZONS:
            base = subset[f"T+{horizon}"]
            adjusted = subset[f"adjusted_T+{horizon}"]
            active_h = active[f"T+{horizon}"]
            row[f"T+{horizon}_return"] = float(base.mean())
            row[f"T+{horizon}_win_rate"] = float(base.gt(0).mean())
            row[f"T+{horizon}_adjusted_return"] = float(adjusted.mean())
            row[f"T+{horizon}_adjusted_win_rate"] = float(adjusted.gt(0).mean())
            row[f"T+{horizon}_active_return"] = float(active_h.mean())
            row[f"T+{horizon}_active_win_rate"] = float(active_h.gt(0).mean())
        rows.append(row)
    return pd.DataFrame(rows)


def make_long_summary(summary: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for _, row in summary.iterrows():
        for horizon in HORIZONS:
            rows.append(
                {
                    "period": row["period"],
                    "horizon": f"T+{horizon}",
                    "events": row["events"],
                    "veto_events": row["veto_events"],
                    "baseline_return": row[f"T+{horizon}_return"],
                    "baseline_win_rate": row[f"T+{horizon}_win_rate"],
                    "adjusted_return": row[f"T+{horizon}_adjusted_return"],
                    "adjusted_win_rate": row[f"T+{horizon}_adjusted_win_rate"],
                    "return_delta": row[f"T+{horizon}_adjusted_return"] - row[f"T+{horizon}_return"],
                    "win_rate_delta": row[f"T+{horizon}_adjusted_win_rate"] - row[f"T+{horizon}_win_rate"],
                }
            )
    return pd.DataFrame(rows)


def make_counts(
    raw_signal: pd.DataFrame,
    gated_signal: pd.DataFrame,
    events: pd.DataFrame,
    latest_signal_date: pd.Timestamp,
    start_date: pd.Timestamp,
    path_excluded: pd.DataFrame | None = None,
) -> pd.DataFrame:
    raw = raw_signal.loc[start_date:latest_signal_date]
    gated = gated_signal.loc[start_date:latest_signal_date]
    path_excluded_count = 0 if path_excluded is None else int(len(path_excluded))
    rows = [
            {
                "stage": "raw signal",
                "events": int(raw.sum().sum()),
                "signal_days": int(raw.any(axis=1).sum()),
                "symbols": int(raw.any(axis=0).sum()),
                "excluded_from_previous": 0,
            },
            {
                "stage": "KOSPI200 + tradable",
                "events": int(gated.sum().sum()),
                "signal_days": int(gated.any(axis=1).sum()),
                "symbols": int(gated.any(axis=0).sum()),
                "excluded_from_previous": int(raw.sum().sum() - gated.sum().sum()),
            },
            {
                "stage": "same-symbol T+1~T+5 dedup",
                "events": int(len(events) + path_excluded_count),
                "signal_days": int(events["signal_date"].nunique()),
                "symbols": int(events["symbol"].nunique()),
                "excluded_from_previous": int(gated.sum().sum() - len(events) - path_excluded_count),
            },
            {
                "stage": "path-normal tradable T+1~T+5",
                "events": int(len(events)),
                "signal_days": int(events["signal_date"].nunique()),
                "symbols": int(events["symbol"].nunique()),
                "excluded_from_previous": path_excluded_count,
            },
            {
                "stage": "adjusted entry veto",
                "events": int(len(events) - events["gap_over_signal_high"].sum()),
                "signal_days": int(events.loc[~events["gap_over_signal_high"], "signal_date"].nunique()),
                "symbols": int(events.loc[~events["gap_over_signal_high"], "symbol"].nunique()),
                "excluded_from_previous": int(events["gap_over_signal_high"].sum()),
            },
    ]
    return pd.DataFrame(rows)


def make_figures(
    summary: pd.DataFrame,
    long_summary: pd.DataFrame,
    counts: pd.DataFrame,
    events: pd.DataFrame,
    open_: pd.DataFrame,
    high: pd.DataFrame,
    low: pd.DataFrame,
    close: pd.DataFrame,
) -> dict[str, Path]:
    plt.rcParams["font.family"] = ["Malgun Gothic", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False
    plt.rcParams["figure.facecolor"] = "white"
    plt.rcParams["axes.facecolor"] = "white"
    figures: dict[str, Path] = {}
    stale_best_exit = FIG_DIR / "adjusted_best_exit_timing.png"
    if stale_best_exit.exists():
        stale_best_exit.unlink()

    n_periods = len(summary)
    fig, axes = plt.subplots(1, n_periods, figsize=(7.4 if n_periods == 1 else 4.1 * n_periods, 3.9), sharey=True)
    if n_periods == 1:
        axes = [axes]
    for ax, (_, row) in zip(axes, summary.iterrows(), strict=True):
        x = np.arange(len(HORIZONS))
        baseline = [row[f"T+{h}_return"] * 100 for h in HORIZONS]
        adjusted = [row[f"T+{h}_adjusted_return"] * 100 for h in HORIZONS]
        ax.bar(x, baseline, width=0.64, label="기존", color=GRAY_BAR, edgecolor="none", zorder=2)
        ax.plot(x, adjusted, label="개선", color=TEAL, marker="o", linewidth=2.6, markersize=5.5, zorder=4)
        ax.fill_between(x, adjusted, baseline, color=TEAL, alpha=0.10, zorder=1)
        ax.axhline(0, color=LINE, linewidth=1.0)
        ax.set_title(str(row["period"]), loc="left", fontweight="bold", color=NAVY, pad=10)
        ax.set_xticks(x, [f"T+{h}" for h in HORIZONS])
        ax.set_ylabel("평균 숏 수익률(%)", color=SLATE)
        style_axis(ax)
        y_low = min(min(baseline), min(adjusted), 0)
        y_high = max(max(baseline), max(adjusted), 0)
        pad = max((y_high - y_low) * 0.22, 0.22)
        ax.set_ylim(y_low - pad, y_high + pad)
        for i, value in enumerate(adjusted):
            offset = pad * 0.16 if value >= 0 else -pad * 0.16
            ax.text(i, value + offset, f"{value:+.2f}%", ha="center", va="bottom" if value >= 0 else "top", fontsize=8.5, color=TEAL, fontweight="bold")
    axes[0].legend(frameon=False, fontsize=9, loc="upper right")
    add_figure_title(fig, "청산 시점별 평균 수익률", "회색 막대는 기존 전략, 청록 라인은 신호일 고가 stop 청산 반영 후 수익률")
    path = FIG_DIR / "adjusted_return_by_exit_horizon.png"
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    figures["return_by_exit"] = path

    figures["performance_matrix"] = make_performance_matrix_figure(long_summary)
    figures["comparison_matrix"] = make_comparison_matrix_figure(long_summary)
    figures["candle_examples"] = make_candle_examples_figure(events, open_, high, low, close)
    return figures


def style_axis(ax, *, axis: str = "y") -> None:
    ax.grid(axis=axis, color=LINE, linewidth=0.8, alpha=0.65)
    ax.tick_params(colors=SLATE, labelsize=8.5)
    for spine in ("top", "right", "left"):
        ax.spines[spine].set_visible(False)
    ax.spines["bottom"].set_color(LINE)


def add_figure_title(fig, title: str, subtitle: str) -> None:
    fig.suptitle(title, x=0.02, y=1.04, ha="left", fontweight="bold", color=NAVY, fontsize=14)
    fig.text(0.02, 0.965, subtitle, ha="left", color=SLATE, fontsize=9)
    fig.tight_layout(rect=[0, 0, 1, 0.92])


def make_performance_matrix_figure(long_summary: pd.DataFrame) -> Path:
    periods = list(dict.fromkeys(long_summary["period"]))
    cell_text = []
    cell_colors = []
    header = ["기간", "이벤트", *[f"T+{h}" for h in HORIZONS]]
    for period in periods:
        subset = long_summary[long_summary["period"].eq(period)]
        first = subset.iloc[0]
        row = [period, f"{int(first['events']):,}건"]
        colors_row = ["#FFFFFF", "#FFFFFF"]
        for _, item in subset.iterrows():
            row.append(f"{item['baseline_return']*100:+.2f}%\n승률 {item['baseline_win_rate']*100:.1f}%")
            colors_row.append(return_color(item["baseline_return"]))
        cell_text.append(row)
        cell_colors.append(colors_row)

    fig, ax = plt.subplots(figsize=(10.2, 2.75))
    ax.axis("off")
    ax.text(0.0, 1.16, "기존 전략 Event Study 테이블", transform=ax.transAxes, fontsize=14, fontweight="bold", color=NAVY)
    ax.text(0.0, 1.05, "각 셀은 평균 숏 수익률과 승률을 함께 표시", transform=ax.transAxes, fontsize=9, color=SLATE)
    table = ax.table(
        cellText=cell_text,
        colLabels=header,
        cellLoc="center",
        colLoc="center",
        cellColours=cell_colors,
        colColours=[NAVY] * len(header),
        bbox=[0.0, 0.0, 1.0, 0.92],
    )
    style_matplotlib_table(table, header_color=NAVY)
    path = FIG_DIR / "performance_matrix_table.png"
    fig.savefig(path, dpi=190, bbox_inches="tight")
    plt.close(fig)
    return path


def make_comparison_matrix_figure(long_summary: pd.DataFrame) -> Path:
    rows = []
    colors = []
    for _, item in long_summary.iterrows():
        rows.append(
            [
                item["period"],
                item["horizon"],
                f"{item['baseline_return']*100:+.2f}%",
                f"{item['adjusted_return']*100:+.2f}%",
                f"{item['return_delta']*100:+.2f}%",
                f"{item['baseline_win_rate']*100:.1f}%",
                f"{item['adjusted_win_rate']*100:.1f}%",
            ]
        )
        colors.append(["#FFFFFF", "#FFFFFF", "#FFFFFF", return_color(item["adjusted_return"]), delta_color(item["return_delta"]), "#FFFFFF", "#FFFFFF"])

    fig, ax = plt.subplots(figsize=(10.3, 4.6))
    ax.axis("off")
    ax.text(0.0, 1.10, "신호일 고가 stop 청산 적용 전후 비교", transform=ax.transAxes, fontsize=14, fontweight="bold", color=NAVY)
    ax.text(0.0, 1.03, "미진입 이벤트는 0, 진입 후 갭상승은 시가 청산, 장중 터치는 신호일 고가 청산으로 반영", transform=ax.transAxes, fontsize=9, color=SLATE)
    table = ax.table(
        cellText=rows,
        colLabels=["기간", "청산", "기존 수익률", "개선 수익률", "차이", "기존 승률", "개선 승률"],
        cellLoc="center",
        colLoc="center",
        cellColours=colors,
        colColours=[NAVY] * 7,
        bbox=[0.0, 0.0, 1.0, 0.96],
    )
    style_matplotlib_table(table, header_color=NAVY, font_size=8.2)
    path = FIG_DIR / "comparison_matrix_table.png"
    fig.savefig(path, dpi=190, bbox_inches="tight")
    plt.close(fig)
    return path


def make_candle_examples_figure(events: pd.DataFrame, open_: pd.DataFrame, high: pd.DataFrame, low: pd.DataFrame, close: pd.DataFrame) -> Path:
    examples = select_case_examples(events)
    fig, axes = plt.subplots(2, 3, figsize=(12.8, 7.4), sharey=False)
    flat_axes = list(axes.ravel())
    for ax, (_, event) in zip(flat_axes, examples.iterrows(), strict=False):
        plot_candle_example(ax, event, open_, high, low, close)
    for ax in flat_axes[len(examples) :]:
        ax.axis("off")
    fig.suptitle("실제 사례: 수익 실현 / stop 청산 / 미진입", x=0.02, y=1.02, ha="left", fontsize=14, color=NAVY, fontweight="bold")
    fig.text(0.02, 0.975, "파란 점은 진입, 보라 점은 stop 청산(갭상승 시가 포함), 붉은 점선은 고정 청산 기준", ha="left", color=SLATE, fontsize=9)
    fig.tight_layout(rect=[0, 0, 1, 0.94], h_pad=2.0, w_pad=1.6)
    path = FIG_DIR / "entry_and_no_entry_candle_examples.png"
    fig.savefig(path, dpi=190, bbox_inches="tight")
    plt.close(fig)
    return path


def select_case_examples(events: pd.DataFrame) -> pd.DataFrame:
    preferred_profit = [
        ("2026-05-15", "A004020"),
        ("2026-03-06", "A079550"),
        ("2026-02-27", "A307950"),
    ]
    preferred_intraday_stop = [
        ("2026-05-18", "A402340"),
    ]
    preferred_gap_stop = [
        ("2026-03-17", "A009830"),
    ]
    preferred_not_entered = [
        ("2026-03-04", "A000660"),
    ]
    entered_touch = events[events["entry_status"].eq("entered") & events["first_prior_high_touch_date"].notna()].copy()
    profitable = events[
        events["entry_status"].eq("entered")
        & events["first_prior_high_touch_date"].isna()
        & events["adjusted_T+5"].gt(0)
    ].copy()
    gap_open_stop = entered_touch[
        entered_touch[[f"adjusted_exit_reason_T{horizon}" for horizon in HORIZONS]].eq("gap_open_stop").any(axis=1)
    ].copy()
    not_entered = events[events["entry_status"].eq("not_entered")].copy()
    selected: list[pd.Series] = []

    add_preferred_examples(selected, profitable, preferred_profit)
    add_preferred_examples(selected, entered_touch, preferred_intraday_stop)
    add_preferred_examples(selected, gap_open_stop, preferred_gap_stop)
    for date, symbol in preferred_not_entered:
        match = not_entered[not_entered["signal_date"].eq(pd.Timestamp(date)) & not_entered["symbol"].eq(symbol)]
        if not match.empty:
            selected.append(match.iloc[0])
    if len(selected) < 6:
        fallback_profit = profitable.sort_values(["adjusted_T+5", "signal_date"], ascending=[False, False])
        for _, row in fallback_profit.iterrows():
            if len(selected) >= 6:
                break
            if not any(row["symbol"] == item["symbol"] for item in selected):
                selected.append(row)
    if len(selected) < 6:
        fallback_entered = entered_touch.sort_values("signal_date", ascending=False)
        for _, row in fallback_entered.iterrows():
            if len(selected) >= 6:
                break
            if not any(pd.Timestamp(row["signal_date"]) == pd.Timestamp(item["signal_date"]) and row["symbol"] == item["symbol"] for item in selected):
                selected.append(row)
    return pd.DataFrame(selected[:6])


def add_preferred_examples(selected: list[pd.Series], events: pd.DataFrame, preferred: list[tuple[str, str]]) -> None:
    for date, symbol in preferred:
        match = events[events["signal_date"].eq(pd.Timestamp(date)) & events["symbol"].eq(symbol)]
        if match.empty:
            continue
        row = match.iloc[0]
        if not any(pd.Timestamp(row["signal_date"]) == pd.Timestamp(item["signal_date"]) and row["symbol"] == item["symbol"] for item in selected):
            selected.append(row)


def plot_candle_example(ax, event: pd.Series, open_: pd.DataFrame, high: pd.DataFrame, low: pd.DataFrame, close: pd.DataFrame) -> None:
    signal_date = pd.Timestamp(event["signal_date"])
    symbol = str(event["symbol"])
    index = open_.index
    pos = index.get_loc(signal_date)
    start = max(0, pos - 5)
    end = min(len(index), pos + 6)
    dates = index[start:end]
    x = np.arange(len(dates))
    o = open_.loc[dates, symbol].astype(float)
    h = high.loc[dates, symbol].astype(float)
    low_series = low.loc[dates, symbol].astype(float)
    c = close.loc[dates, symbol].astype(float)
    signal_x = int(np.where(dates == signal_date)[0][0])
    entry_date = pd.Timestamp(event["entry_date"])
    entry_x = int(np.where(dates == entry_date)[0][0])
    entered = str(event.get("entry_status", "entered")) == "entered"
    touch_date = pd.Timestamp(event["first_prior_high_touch_date"]) if pd.notna(event.get("first_prior_high_touch_date")) else pd.NaT

    ax.axvspan(signal_x - 0.45, signal_x + 0.45, color=GOLD, alpha=0.17, zorder=0)
    ax.axhline(float(event["signal_high"]), color=CORAL, linestyle=":", linewidth=1.4)
    if entered:
        ax.axvline(entry_x, color=TEAL, linestyle="--", linewidth=1.2, alpha=0.85)
    else:
        ax.axvline(entry_x, color=CORAL, linestyle="--", linewidth=1.2, alpha=0.85)
    for i in x:
        color = CORAL if c.iloc[i] >= o.iloc[i] else TEAL
        ax.vlines(i, low_series.iloc[i], h.iloc[i], color=color, linewidth=1.15, alpha=0.95)
        bottom = min(o.iloc[i], c.iloc[i])
        height = abs(c.iloc[i] - o.iloc[i])
        height = max(height, max(h.max() - low_series.min(), 1) * 0.006)
        ax.add_patch(Rectangle((i - 0.28, bottom), 0.56, height, facecolor=color, edgecolor=color, alpha=0.85))
    ax.scatter([entry_x], [float(event["entry_open"])], s=48, color=NAVY if entered else CORAL, edgecolor="white", linewidth=1.2, zorder=5)
    if entered and pd.notna(touch_date) and touch_date in set(dates):
        touch_x = int(np.where(dates == touch_date)[0][0])
        touch_price = float(event["first_prior_high_touch_price"])
        ax.scatter([touch_x], [touch_price], s=58, color="#6A3D9A", edgecolor="white", linewidth=1.2, zorder=6)
        ax.axvline(touch_x, color="#6A3D9A", linestyle=":", linewidth=1.1, alpha=0.8)
    label_name = str(event.get("name") or "").strip()
    label = f"{label_name} {symbol}" if label_name else symbol
    ax.set_title(f"{label}\n{signal_date:%Y-%m-%d}", loc="left", fontsize=10.5, color=NAVY, fontweight="bold")
    ax.text(0.03, 0.92, "T 신호", transform=ax.transAxes, fontsize=8.5, color=GOLD, fontweight="bold")
    if entered and pd.notna(touch_date):
        exit_reason = first_adjusted_exit_reason(event)
        if exit_reason == "gap_open_stop":
            note = f"진입 후 청산\n{touch_date:%m/%d} 갭상승 시가"
        else:
            note = f"진입 후 청산\n{touch_date:%m/%d} 신호일 고가"
        note_color = "#6A3D9A"
    elif entered:
        t5_return = float(event.get("adjusted_T+5", np.nan))
        if pd.notna(t5_return) and t5_return > 0:
            note = f"수익 사례\nT+5 {t5_return * 100:+.1f}%"
            note_color = TEAL
        else:
            note = "진입 유지\nhorizon 종가 청산"
            note_color = TEAL
    else:
        note = "T+1 시가 > T 고가\n미진입"
        note_color = CORAL
    ax.text(0.56, 0.08, note, transform=ax.transAxes, fontsize=7.8, color=note_color, fontweight="bold")
    tick_labels = [d.strftime("%m/%d") if i in {0, signal_x, entry_x, len(dates) - 1} else "" for i, d in enumerate(dates)]
    ax.set_xticks(x, tick_labels, fontsize=7.6)
    ax.tick_params(axis="y", labelsize=7.6, colors=SLATE)
    ax.grid(axis="y", color=LINE, linewidth=0.7, alpha=0.6)
    for spine in ("top", "right", "left"):
        ax.spines[spine].set_visible(False)
    ax.spines["bottom"].set_color(LINE)


def first_adjusted_exit_reason(event: pd.Series) -> str:
    for horizon in HORIZONS:
        if bool(event.get(f"prior_high_touch_T{horizon}", False)):
            return str(event.get(f"adjusted_exit_reason_T{horizon}", "prior_high_touch"))
    return ""


def return_color(value: float) -> str:
    if value >= 0.01:
        return "#D8F3F0"
    if value >= 0.003:
        return "#EAF8F7"
    if value >= 0:
        return "#F4FBFA"
    return "#FDE7E4"


def delta_color(value: float) -> str:
    if value >= 0:
        return "#D8F3F0"
    if value > -0.001:
        return "#F7EFEA"
    return "#FDE2DF"


def style_matplotlib_table(table, *, header_color: str, font_size: float = 8.6) -> None:
    table.auto_set_font_size(False)
    table.set_fontsize(font_size)
    for (row, _col), cell in table.get_celld().items():
        cell.set_edgecolor("#E2E8F0")
        cell.set_linewidth(0.65)
        if row == 0:
            cell.set_text_props(color="white", weight="bold")
            cell.set_height(0.16)
        else:
            cell.set_text_props(color=INK)
            cell.set_height(0.18)


def write_workbook(data: ReportData) -> None:
    path = OUT_DIR / f"{REPORT_STEM}.xlsx"
    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        data.counts.to_excel(writer, sheet_name="COUNTS", index=False)
        data.summary.to_excel(writer, sheet_name="PERIOD_SUMMARY", index=False)
        data.long_summary.to_excel(writer, sheet_name="EVENT_STUDY_COMPARISON", index=False)
        data.events.to_excel(writer, sheet_name="EVENTS", index=False)
        data.path_excluded.to_excel(writer, sheet_name="PATH_NORMAL_EXCLUDED", index=False)


def build_docx(data: ReportData, path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    configure_styles(doc)
    set_footer(section)

    add_title(doc, "KOSPI200 유성형 매도 전략")
    add_subtitle(doc, "Event Study 기반 전략 보고서")

    add_heading(doc, "전략 개요", level=1)
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["대상", "KOSPI200 구성 종목"],
            ["신호", "유성형: 큰 장중 변동, 긴 윗꼬리, 63거래일 급등, 거래량 상위 10% 동시 충족"],
            ["진입", "신호일 T 다음 영업일(T+1) 시가에 숏 진입"],
            ["청산", "T+1~T+5 종가 청산 후보와 신호일 고가 터치 청산을 비교"],
            ["주기", "모든 영업일"],
        ],
        widths=[1.25, 5.05],
    )

    add_heading(doc, "전략 설명", level=1)
    add_table(
        doc,
        ["조건", "판단 기준", "전략상 의의"],
        [
            ["장중 변동성", "(고가 - 저가) / 전일 종가 > 10%", "과열 또는 투기적 거래가 충분히 커진 날만 선별"],
            ["윗꼬리", "고가 - max(시가, 종가) > |종가 - 시가|", "고점 부근 매물 출회와 장중 되밀림 확인"],
            ["선행 급등", "고가 / 63거래일 전 종가 > 1.5배", "단순 변동이 아니라 누적 과열 이후의 유성형만 채택"],
            ["거래량 확인", "252거래일 거래량 백분위 > 90%", "수급 집중이 동반된 이벤트로 제한"],
            ["동일 종목 압축", "채택 이벤트 후 T+1~T+5 같은 종목 후속 신호 제거", "같은 급등 국면의 중복 카운트를 줄여 이벤트 독립성 개선"],
        ],
        widths=[1.35, 2.55, 2.4],
    )

    add_heading(doc, "기타 요건", level=1)
    add_table(
        doc,
        ["요건", "적용 방식"],
        [
            ["분석 기간", f"{fmt_date(ANALYSIS_START_DATE)}~{fmt_date(ANALYSIS_END_DATE)}"],
            ["거래 정지 구분", f"신호일 T와 T+1~T+5 경로의 거래정지 필드가 '{NORMAL_TRADING_LABEL}'인 경우만 사용"],
            ["KOSPI200 구성 종목 구분", "신호일 T 기준 KOSPI200 편입 플래그가 1인 종목만 사용"],
            ["중복 신호 처리", "동일 종목의 채택 이벤트 이후 T+1~T+5 신호는 제외하고 T+6부터 다시 허용"],
            ["수익률 산식", "(T+1 시가 - T+h 종가) / T+1 시가; 숏 관점 수익률"],
            ["평균 방식", "포트폴리오 가중 평균이 아닌 이벤트 1건 단순평균"],
            ["실거래 미반영", "대차 가능성, 공매도 비용, 호가 충격, 세금, 슬리피지는 제외"],
        ],
        widths=[1.55, 4.75],
    )

    add_heading(doc, "전략 성과", level=1)
    add_para(
        doc,
        "아래 표는 2016-01-01~2026-05-27 단일 분석 구간의 기존 유성형 매도 신호 event study 결과다. 승률은 해당 horizon 수익률이 0보다 큰 이벤트 비율이다.",
    )
    doc.add_picture(str(data.figures["performance_matrix"]), width=Inches(6.0))
    last_para(doc).alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "경제적 함의", level=1)
    add_table(
        doc,
        ["원인", "이유"],
        [
            ["과열 후 단기 되돌림", "63거래일 1.5배 이상 급등 뒤 대량 거래와 윗꼬리가 겹치면 단기 차익실현 압력이 커진다."],
            ["효과의 짧은 반감기", "2016년 이후 표본에서는 T+1 수익률이 가장 안정적이고, 보유 기간이 길어질수록 평균이 약해진다."],
            ["손실 꼬리 존재", "일부 급등 지속 종목이 평균과 승률을 훼손하므로 무조건 보유보다 빠른 확인/청산 규칙이 필요하다."],
            ["실행 제약", "KOSPI200 유동성은 완화 요인이지만 공매도 가능 수량, 대차 비용, 갭 상승은 별도 관리 대상이다."],
        ],
        widths=[1.55, 4.75],
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "유성형 매도 전략 개선 연구", level=1)
    add_heading(doc, "1. 청산 필터 추가", level=2)
    add_para(
        doc,
        "개선안은 T+1 시가가 신호일 T의 고가를 상향 돌파하면 미진입으로 처리하고, 진입한 이벤트는 보유 중 신호일 고가를 stop 기준으로 삼아 청산하는 방식이다. T+2 이후 시가가 이미 stop 위에서 시작하면 해당 시가로 청산하고, 장중에만 stop을 터치하면 신호일 고가 가격으로 청산한다.",
    )
    add_table(
        doc,
        ["필터", "판단", "처리"],
        [
            ["T+1 시가 > T 고가", "신호일 고가가 다음날 시가에서 돌파됨", "미진입, 수익률 0"],
            ["T+1 시가 <= T 고가", "신호일 고가가 아직 유효한 저항선", "T+1 시가 진입"],
            ["갭상승 stop", "진입 후 T+2 이후 시가 > 신호일 고가", "해당일 시가에 청산"],
            ["장중 stop 터치", "진입 후 당일 고가 >= 신호일 고가", "신호일 고가에 청산"],
            ["미터치", "horizon까지 신호일 고가 터치 없음", "해당 T+h 종가 청산"],
        ],
        widths=[1.6, 2.6, 2.1],
    )
    add_heading(doc, "경제적 함의", level=3)
    add_table(
        doc,
        ["함의", "청산 필터를 추가한 이유"],
        liquidation_filter_rationale_rows(),
        widths=[1.55, 4.75],
    )
    doc.add_page_break()
    doc.add_picture(str(data.figures["candle_examples"]), width=Inches(6.35))
    last_para(doc).alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "개선안 Event Study 비교", level=2)
    add_para(doc, "개선 전략은 미진입 이벤트를 0 수익률로 두고, 진입 이벤트는 신호일 고가 stop 청산을 적용했다. 갭상승 출발은 시가 청산, 장중 터치는 stop 가격 청산으로 반영해 단순 horizon 종가 청산과 비교한다.")
    doc.add_picture(str(data.figures["comparison_matrix"]), width=Inches(6.0))
    last_para(doc).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(str(data.figures["return_by_exit"]), width=Inches(5.65))
    last_para(doc).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.save(path)


def build_pdf(data: ReportData, path: Path) -> None:
    register_pdf_fonts()
    doc = SimpleDocTemplate(
        str(path),
        pagesize=landscape(letter),
        rightMargin=0.42 * inch,
        leftMargin=0.42 * inch,
        topMargin=0.38 * inch,
        bottomMargin=0.38 * inch,
        title="report_adjusted",
    )
    styles = pdf_styles()
    story: list[object] = []

    story.append(Paragraph("KOSPI200 유성형 매도 전략", styles["TitleKR"]))
    story.append(Paragraph("Event Study 기반 전략 보고서", styles["SubtitleKR"]))

    story.append(Paragraph("전략 개요", styles["H1KR"]))
    story.append(
        pdf_table(
            [["항목", "내용"]]
            + [
                ["대상", "KOSPI200 구성 종목"],
                ["신호", "유성형: 큰 장중 변동, 긴 윗꼬리, 63거래일 급등, 거래량 상위 10% 동시 충족"],
                ["진입", "신호일 T 다음 영업일(T+1) 시가에 숏 진입"],
                ["청산", "T+1~T+5 종가 청산 후보와 신호일 고가 터치 청산을 비교"],
                ["주기", "모든 영업일"],
            ],
            [1.0 * inch, 8.8 * inch],
            styles,
        )
    )

    story.append(Paragraph("전략 설명", styles["H1KR"]))
    story.append(
        pdf_table(
            [["조건", "판단 기준", "전략상 의의"]]
            + [
                ["장중 변동성", "(고가 - 저가) / 전일 종가 > 10%", "과열 또는 투기적 거래가 충분히 커진 날만 선별"],
                ["윗꼬리", "고가 - max(시가, 종가) > |종가 - 시가|", "고점 부근 매물 출회와 장중 되밀림 확인"],
                ["선행 급등", "고가 / 63거래일 전 종가 > 1.5배", "단순 변동이 아니라 누적 과열 이후의 유성형만 채택"],
                ["거래량 확인", "252거래일 거래량 백분위 > 90%", "수급 집중이 동반된 이벤트로 제한"],
                ["동일 종목 압축", "채택 이벤트 후 T+1~T+5 같은 종목 후속 신호 제거", "같은 급등 국면의 중복 카운트를 줄여 이벤트 독립성 개선"],
            ],
            [1.3 * inch, 3.5 * inch, 5.0 * inch],
            styles,
        )
    )

    story.append(Paragraph("기타 요건", styles["H1KR"]))
    story.append(
        pdf_table(
            [["요건", "적용 방식"]]
            + [
                ["분석 기간", f"{fmt_date(ANALYSIS_START_DATE)}~{fmt_date(ANALYSIS_END_DATE)}"],
                ["거래 정지 구분", f"신호일 T와 T+1~T+5 경로의 거래정지 필드가 '{NORMAL_TRADING_LABEL}'인 경우만 사용"],
                ["KOSPI200 구성 종목 구분", "신호일 T 기준 KOSPI200 편입 플래그가 1인 종목만 사용"],
                ["중복 신호 처리", "동일 종목의 채택 이벤트 이후 T+1~T+5 신호는 제외하고 T+6부터 다시 허용"],
                ["수익률 산식", "(T+1 시가 - T+h 종가) / T+1 시가; 숏 관점 수익률"],
                ["평균 방식", "포트폴리오 가중 평균이 아닌 이벤트 1건 단순평균"],
                ["실거래 미반영", "대차 가능성, 공매도 비용, 호가 충격, 세금, 슬리피지는 제외"],
            ],
            [1.45 * inch, 8.35 * inch],
            styles,
        )
    )

    story.append(PageBreak())
    story.append(Paragraph("전략 성과", styles["H1KR"]))
    story.append(Paragraph("2016-01-01~2026-05-27 단일 분석 구간의 기존 유성형 매도 신호 event study 결과다. 승률은 해당 horizon 수익률이 0보다 큰 이벤트 비율이다.", styles["BodyKR"]))
    story.append(pdf_image(data.figures["performance_matrix"], width=8.2 * inch))
    story.append(PageBreak())
    story.append(Paragraph("경제적 함의", styles["H1KR"]))
    story.append(Paragraph("이 전략의 경제적 의미는 지속 하락 추세 예측보다, 급등 이후 단기 수급 과열이 해소되는 구간을 포착하는 데 있다.", styles["BodyKR"]))
    story.append(
        pdf_table(
            [["원인", "이유"]]
            + [
                ["과열 후 단기 되돌림", "63거래일 1.5배 이상 급등 뒤 대량 거래와 윗꼬리가 겹치면 단기 차익실현 압력이 커진다."],
                ["효과의 짧은 반감기", "2016년 이후 표본에서는 T+1 수익률이 가장 안정적이고, 보유 기간이 길어질수록 평균이 약해진다."],
                ["손실 꼬리 존재", "일부 급등 지속 종목이 평균과 승률을 훼손하므로 무조건 보유보다 빠른 확인/청산 규칙이 필요하다."],
                ["실행 제약", "KOSPI200 유동성은 완화 요인이지만 공매도 가능 수량, 대차 비용, 갭 상승은 별도 관리 대상이다."],
            ],
            [1.65 * inch, 8.15 * inch],
            styles,
        )
    )

    story.append(PageBreak())
    story.append(Paragraph("유성형 매도 전략 개선 연구", styles["H1KR"]))
    story.append(Paragraph("1. 청산 필터 추가", styles["H2KR"]))
    story.append(
        Paragraph(
            "개선안은 T+1 시가가 신호일 T의 고가를 상향 돌파하면 미진입으로 처리하고, 진입한 이벤트는 보유 중 신호일 고가를 stop 기준으로 삼아 청산하는 방식이다. T+2 이후 시가가 이미 stop 위에서 시작하면 해당 시가로 청산한다.",
            styles["BodyKR"],
        )
    )
    story.append(
        pdf_table(
            [["필터", "판단", "처리"]]
            + [
                ["T+1 시가 > T 고가", "신호일 고가가 다음날 시가에서 돌파됨", "미진입, 수익률 0"],
                ["T+1 시가 <= T 고가", "신호일 고가가 아직 유효한 저항선", "T+1 시가 진입"],
                ["갭상승 stop", "진입 후 T+2 이후 시가 > 신호일 고가", "해당일 시가에 청산"],
                ["장중 stop 터치", "진입 후 당일 고가 >= 신호일 고가", "신호일 고가에 청산"],
                ["미터치", "horizon까지 신호일 고가 터치 없음", "해당 T+h 종가 청산"],
            ],
            [1.7 * inch, 4.1 * inch, 4.0 * inch],
            styles,
        )
    )
    story.append(Paragraph("경제적 함의", styles["H2KR"]))
    story.append(
        pdf_table(
            [["함의", "청산 필터를 추가한 이유"]] + liquidation_filter_rationale_rows(),
            [1.9 * inch, 7.9 * inch],
            styles,
        )
    )
    story.append(PageBreak())
    story.append(pdf_image(data.figures["candle_examples"], width=9.4 * inch))

    story.append(PageBreak())
    story.append(Paragraph("개선안 Event Study 비교", styles["H2KR"]))
    story.append(Paragraph("개선 전략은 미진입 이벤트를 0 수익률로 두고, 진입 이벤트는 신호일 고가 stop 청산을 적용했다. 갭상승 출발은 시가 청산, 장중 터치는 stop 가격 청산으로 반영했다.", styles["BodyKR"]))
    story.append(pdf_image(data.figures["comparison_matrix"], width=8.3 * inch))
    story.append(Spacer(1, 0.15 * inch))
    story.append(pdf_image(data.figures["return_by_exit"], width=7.6 * inch))

    doc.build(story, onFirstPage=pdf_footer, onLaterPages=pdf_footer)


def register_pdf_fonts() -> None:
    regular = Path("C:/Windows/Fonts/malgun.ttf")
    bold = Path("C:/Windows/Fonts/malgunbd.ttf")
    if regular.exists():
        pdfmetrics.registerFont(TTFont("Malgun", str(regular)))
    if bold.exists():
        pdfmetrics.registerFont(TTFont("Malgun-Bold", str(bold)))


def pdf_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "TitleKR": ParagraphStyle(
            "TitleKR",
            parent=base["Title"],
            fontName="Malgun-Bold",
            fontSize=24,
            leading=29,
            textColor=colors.HexColor("#202B36"),
            alignment=TA_LEFT,
            spaceAfter=5,
        ),
        "SubtitleKR": ParagraphStyle(
            "SubtitleKR",
            parent=base["BodyText"],
            fontName="Malgun",
            fontSize=10,
            leading=13,
            textColor=colors.HexColor("#62707D"),
            spaceAfter=10,
        ),
        "H1KR": ParagraphStyle(
            "H1KR",
            parent=base["Heading1"],
            fontName="Malgun-Bold",
            fontSize=14,
            leading=18,
            textColor=colors.HexColor("#1F4D78"),
            spaceBefore=8,
            spaceAfter=5,
        ),
        "H2KR": ParagraphStyle(
            "H2KR",
            parent=base["Heading2"],
            fontName="Malgun-Bold",
            fontSize=12,
            leading=15,
            textColor=colors.HexColor("#1F4D78"),
            spaceBefore=7,
            spaceAfter=4,
        ),
        "BodyKR": ParagraphStyle(
            "BodyKR",
            parent=base["BodyText"],
            fontName="Malgun",
            fontSize=9.3,
            leading=12.2,
            textColor=colors.HexColor("#202B36"),
            spaceAfter=5,
        ),
        "CellKR": ParagraphStyle(
            "CellKR",
            parent=base["BodyText"],
            fontName="Malgun",
            fontSize=7.7,
            leading=10,
            textColor=colors.HexColor("#202B36"),
            alignment=TA_LEFT,
        ),
        "HeaderCellKR": ParagraphStyle(
            "HeaderCellKR",
            parent=base["BodyText"],
            fontName="Malgun-Bold",
            fontSize=7.8,
            leading=10,
            textColor=colors.HexColor("#202B36"),
            alignment=TA_CENTER,
        ),
        "NoteKR": ParagraphStyle(
            "NoteKR",
            parent=base["BodyText"],
            fontName="Malgun",
            fontSize=8.8,
            leading=11.5,
            textColor=colors.HexColor("#62707D"),
            spaceAfter=4,
        ),
    }


def pdf_table(rows: list[list[object]], widths: list[float], styles: dict[str, ParagraphStyle], *, compact: bool = False) -> Table:
    table_rows = []
    for r_idx, row in enumerate(rows):
        style = styles["HeaderCellKR"] if r_idx == 0 else styles["CellKR"]
        table_rows.append([Paragraph(str(value), style) for value in row])
    table = Table(table_rows, colWidths=widths, repeatRows=1, hAlign="LEFT")
    font_size = 7.2 if compact else 7.7
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#B8C2CC")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 4 if not compact else 2.5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4 if not compact else 2.5),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("FONTSIZE", (0, 0), (-1, -1), font_size),
            ]
        )
    )
    return table


def pdf_image(path: Path, *, width: float) -> Image:
    with PILImage.open(path) as source:
        aspect = source.height / source.width
    img = Image(str(path), width=width, height=width * aspect)
    img.hAlign = "CENTER"
    return img


def note_pdf(text: str, styles: dict[str, ParagraphStyle]) -> Table:
    table = Table([[Paragraph(text, styles["NoteKR"])]], colWidths=[9.8 * inch], hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F7F9FC")),
                ("BOX", (0, 0), (-1, -1), 0.35, colors.HexColor("#D8DEE8")),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )
    return table


def pdf_footer(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFont("Malgun", 7.5)
    canvas.setFillColor(colors.HexColor("#62707D"))
    canvas.drawRightString(10.55 * inch, 0.2 * inch, f"KOSPI200 유성형 매도 전략 | report_adjusted | {doc.page}")
    canvas.restoreState()


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(9.6)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.1
    for style_name, size, before, after in [
        ("Heading 1", 15, 13, 6),
        ("Heading 2", 12.5, 10, 4),
        ("Heading 3", 11, 7, 3),
    ]:
        style = styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = BLUE
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def set_footer(section) -> None:
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("KOSPI200 유성형 매도 전략 | report_adjusted")
    set_run_font(r, size=8.5, color=MUTED)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_run_font(r, size=23, bold=True, color=DARK)


def add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=MUTED)


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.add_run(text)


def add_note(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="D8DEE8")
    set_cell_shading(table.cell(0, 0), "F7F9FC")
    set_cell_margins(table.cell(0, 0), top=90, bottom=90, start=130, end=130)
    p = table.cell(0, 0).paragraphs[0]
    r = p.add_run(text)
    set_run_font(r, size=9, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[object]],
    widths: list[float],
    font_size: float = 8.8,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color=BORDER)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = Inches(widths[i])
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(header))
        set_run_font(r, size=font_size, bold=True, color=DARK)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            cell.width = Inches(widths[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 and len(str(value)) <= 18 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def event_rows(long_summary: pd.DataFrame, *, adjusted: bool) -> list[list[str]]:
    rows = []
    for _, row in long_summary.iterrows():
        prefix = "adjusted_" if adjusted else "baseline_"
        rows.append(
            [
                row["period"],
                row["horizon"],
                f"{int(row['events']):,}",
                pct(row[f"{prefix}return"]),
                pct(row[f"{prefix}win_rate"]),
            ]
        )
    return rows


def comparison_rows(long_summary: pd.DataFrame) -> list[list[str]]:
    rows = []
    for _, row in long_summary.iterrows():
        rows.append(
            [
                row["period"],
                row["horizon"],
                pct(row["baseline_return"]),
                pct(row["adjusted_return"]),
                pct(row["return_delta"], signed=True),
                pct(row["baseline_win_rate"]),
                pct(row["adjusted_win_rate"]),
            ]
        )
    return rows


def liquidation_filter_rationale_rows() -> list[list[str]]:
    return [
        ["저항선 무효화", "T+1 시가가 신호일 고가를 넘으면 윗꼬리 저항 가정이 이미 깨졌으므로 진입을 피한다."],
        ["손실 꼬리 축소", "진입 후 신호일 고가를 회복하면 급등 지속 구간으로 판단하고 손실 확대 전에 청산한다."],
        ["체결 현실성", "T+2 이후 stop 위에서 갭상승 출발하면 stop 가격이 아니라 해당일 시가 청산으로 보수적으로 반영한다."],
        ["보유기간 관리", "유성형 신호의 평균 우위가 짧게 나타나므로 실패 신호를 빠르게 분리해 장기 보유 리스크를 줄인다."],
    ]


def pct(value: object, *, signed: bool = False) -> str:
    number = float(value) * 100.0
    return f"{number:+.2f}%" if signed else f"{number:.2f}%"


def fmt_date(value: object) -> str:
    return pd.Timestamp(value).strftime("%Y-%m-%d")


def set_run_font(run, *, size: float, color: RGBColor, bold: bool | None = None) -> None:
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Malgun Gothic")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Malgun Gothic")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_table_borders(table, color: str = BORDER) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top: int = 75, bottom: int = 75, start: int = 100, end: int = 100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def last_para(doc: Document):
    return doc.paragraphs[-1]


if __name__ == "__main__":
    main()
