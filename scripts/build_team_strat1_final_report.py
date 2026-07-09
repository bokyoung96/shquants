from __future__ import annotations

# ruff: noqa: E402

from dataclasses import dataclass
from pathlib import Path
import shutil
import sys
import tempfile

import fitz
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from reportlab.lib import colors
from reportlab.lib.pagesizes import landscape, letter
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from scripts import build_team_strat1_adjusted_report as adjusted
from scripts import build_team_strat1_lower_wick_25pct_report as lower25


OUT_DIR = adjusted.OUT_DIR
REPORT_STEM = "report_final"


@dataclass(frozen=True)
class FinalReportData:
    adjusted: adjusted.ReportData
    lower25: lower25.LowerWick25ReportData
    final_summary: pd.DataFrame


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    data = build_report_data()
    write_workbook(data)
    docx_path = OUT_DIR / f"{REPORT_STEM}.docx"
    pdf_path = OUT_DIR / f"{REPORT_STEM}.pdf"
    build_docx(data, docx_path)
    build_pdf(data, pdf_path)
    print(f"wrote {docx_path}")
    print(f"wrote {pdf_path}")
    print(f"wrote {OUT_DIR / f'{REPORT_STEM}.xlsx'}")


def build_report_data() -> FinalReportData:
    adjusted_data = adjusted.build_report_data()
    lower25_data = lower25.build_report_data()
    return FinalReportData(
        adjusted=adjusted_data,
        lower25=lower25_data,
        final_summary=make_final_summary(adjusted_data, lower25_data),
    )


def make_final_summary(
    adjusted_data: adjusted.ReportData,
    lower25_data: lower25.LowerWick25ReportData,
) -> pd.DataFrame:
    stop_row = adjusted_data.summary.iloc[0]
    lower_es = lower25_data.event_study.set_index("horizon")
    rows: list[dict[str, object]] = []
    for horizon in adjusted.HORIZONS:
        rows.append(
            {
                "horizon": f"T+{horizon}",
                "events": int(stop_row["events"]),
                "baseline_return": float(stop_row[f"T+{horizon}_return"]),
                "lower_wick_25_return": float(lower_es.at[f"T+{horizon}", "lower_wick_25_return"]),
                "lower25_delta_vs_baseline_bp": float(lower_es.at[f"T+{horizon}", "delta_vs_baseline"]) * 10000,
            }
        )
    return pd.DataFrame(rows)


def write_workbook(data: FinalReportData) -> None:
    with pd.ExcelWriter(OUT_DIR / f"{REPORT_STEM}.xlsx", engine="openpyxl") as writer:
        data.final_summary.to_excel(writer, sheet_name="FINAL_SUMMARY", index=False)
        data.adjusted.long_summary.to_excel(writer, sheet_name="REPORT_ADJUSTED", index=False)
        data.lower25.event_study.to_excel(writer, sheet_name="LOWER_WICK_25", index=False)
        data.lower25.cases.to_excel(writer, sheet_name="LOWER_WICK_CASES", index=False)


def build_docx(data: FinalReportData, path: Path) -> None:
    source = OUT_DIR / "report_adjusted.docx"
    if not source.exists():
        source = path.with_name("_tmp_report_adjusted.docx")
        adjusted.build_docx(data.adjusted, source)

    if path.exists():
        path.unlink()
    shutil.copy2(source, path)
    doc = Document(path)
    doc.add_section(WD_SECTION.NEW_PAGE)
    append_lower_wick_docx(doc, data)
    doc.save(path)


def append_lower_wick_docx(doc: Document, data: FinalReportData) -> None:
    adjusted.add_heading(doc, "아래꼬리 25% 미진입 후보", level=1)
    adjusted.add_table(doc, ["항목", "내용"], lower25.condition_rows(data.lower25.events), widths=[1.45, 4.85], font_size=8.4)
    adjusted.add_heading(doc, "경제적 함의", level=2)
    adjusted.add_table(doc, ["함의", "설명"], lower25.lower_wick_rationale_rows(), widths=[1.45, 4.85], font_size=8.3)
    adjusted.add_heading(doc, "사례", level=2)
    doc.add_picture(str(data.lower25.figures["cases"]), width=Inches(6.25))
    adjusted.last_para(doc).alignment = WD_ALIGN_PARAGRAPH.CENTER
    adjusted.add_heading(doc, "Event Study 비교", level=2)
    doc.add_picture(str(data.lower25.figures["event_study_matrix"]), width=Inches(6.05))
    adjusted.last_para(doc).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(str(data.lower25.figures["event_study"]), width=Inches(5.75))
    adjusted.last_para(doc).alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_pdf(data: FinalReportData, path: Path) -> None:
    source = OUT_DIR / "report_adjusted.pdf"
    if not source.exists():
        source = path.with_name("_tmp_report_adjusted.pdf")
        adjusted.build_pdf(data.adjusted, source)

    with tempfile.TemporaryDirectory() as tmpdir:
        appendix = Path(tmpdir) / "lower_wick_appendix.pdf"
        build_appendix_pdf(data, appendix)
        merged = fitz.open()
        with fitz.open(source) as base_pdf:
            merged.insert_pdf(base_pdf)
        with fitz.open(appendix) as appendix_pdf:
            merged.insert_pdf(appendix_pdf)
        merged.save(path)
        merged.close()


def build_appendix_pdf(data: FinalReportData, path: Path) -> None:
    adjusted.register_pdf_fonts()
    styles = adjusted.pdf_styles()
    doc = SimpleDocTemplate(
        str(path),
        pagesize=landscape(letter),
        rightMargin=0.46 * inch,
        leftMargin=0.46 * inch,
        topMargin=0.40 * inch,
        bottomMargin=0.38 * inch,
        title="report_final_appendix",
    )
    story: list[object] = [
        Paragraph("아래꼬리 25% 미진입 후보", styles["H1KR"]),
        adjusted.pdf_table([["항목", "내용"]] + lower25.condition_rows(data.lower25.events), [1.45 * inch, 8.35 * inch], styles),
        Paragraph("경제적 함의", styles["H2KR"]),
        adjusted.pdf_table([["함의", "설명"]] + lower25.lower_wick_rationale_rows(), [1.8 * inch, 8.0 * inch], styles),
        PageBreak(),
        Paragraph("사례", styles["H2KR"]),
        adjusted.pdf_image(data.lower25.figures["cases"], width=9.15 * inch),
        PageBreak(),
        Paragraph("Event Study 비교", styles["H2KR"]),
        adjusted.pdf_image(data.lower25.figures["event_study_matrix"], width=8.3 * inch),
        Spacer(1, 0.10 * inch),
        adjusted.pdf_image(data.lower25.figures["event_study"], width=7.1 * inch),
    ]
    doc.build(story, onFirstPage=pdf_footer, onLaterPages=pdf_footer)


def pdf_footer(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFont("Malgun", 7.5)
    canvas.setFillColor(colors.HexColor("#62707D"))
    canvas.drawRightString(10.55 * inch, 0.2 * inch, f"KOSPI200 변동성 매도 전략 | report_final appendix | {doc.page}")
    canvas.restoreState()


if __name__ == "__main__":
    main()
