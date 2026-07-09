from __future__ import annotations

# ruff: noqa: E402

from dataclasses import dataclass
from pathlib import Path
import sys

import matplotlib.pyplot as plt
from matplotlib.patches import Rectangle
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import landscape, letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from scripts import build_team_strat1_adjusted_report as base


OUT_DIR = base.OUT_DIR
FIG_DIR = OUT_DIR / "lower_wick_25pct_figures"
REPORT_STEM = "report_lower_wick_25pct"
THRESHOLD = 0.25


@dataclass(frozen=True)
class LowerWick25ReportData:
    events: pd.DataFrame
    event_study: pd.DataFrame
    cases: pd.DataFrame
    figures: dict[str, Path]


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    data = build_report_data()
    write_workbook(data)
    docx_path = OUT_DIR / f"{REPORT_STEM}.docx"
    pdf_path = OUT_DIR / f"{REPORT_STEM}.pdf"
    build_docx(data, docx_path)
    build_pdf(data, pdf_path)
    print(f"wrote {docx_path}")
    print(f"wrote {pdf_path}")
    print(f"wrote {OUT_DIR / f'{REPORT_STEM}.xlsx'}")


def build_report_data() -> LowerWick25ReportData:
    adjusted = base.build_report_data()
    open_ = base.read_price("qw_adj_o")
    high = base.read_price("qw_adj_h").reindex_like(open_)
    low = base.read_price("qw_adj_l").reindex_like(open_)
    close = base.read_price("qw_adj_c").reindex_like(open_)

    events = add_lower_wick_metrics(adjusted.events.copy(), open_, high, low, close)
    events["lower_wick_25_filter"] = events["lower_wick_ratio"].ge(THRESHOLD)
    events["combined_no_entry_25"] = events["lower_wick_25_filter"]
    for horizon in base.HORIZONS:
        events[f"lower_wick_25_T+{horizon}"] = events[f"T+{horizon}"].mask(
            events["lower_wick_25_filter"],
            0.0,
        )

    event_study = summarize_event_study(events)
    cases = select_case_examples(events)
    figures = {
        "event_study_matrix": make_event_study_matrix_figure(event_study),
        "event_study": make_event_study_figure(event_study),
        "cases": make_case_figure(cases, open_, high, low, close),
    }
    return LowerWick25ReportData(events=events, event_study=event_study, cases=cases, figures=figures)


def add_lower_wick_metrics(
    events: pd.DataFrame,
    open_: pd.DataFrame,
    high: pd.DataFrame,
    low: pd.DataFrame,
    close: pd.DataFrame,
) -> pd.DataFrame:
    rows: list[dict[str, float]] = []
    for _, event in events.iterrows():
        date = pd.Timestamp(event["signal_date"])
        symbol = str(event["symbol"])
        open_price = float(open_.at[date, symbol])
        high_price = float(high.at[date, symbol])
        low_price = float(low.at[date, symbol])
        close_price = float(close.at[date, symbol])
        day_range = high_price - low_price
        body_low = min(open_price, close_price)
        lower_wick = max(body_low - low_price, 0.0)
        rows.append(
            {
                "signal_open": open_price,
                "signal_low": low_price,
                "signal_close": close_price,
                "signal_range": day_range,
                "lower_wick": lower_wick,
                "lower_wick_ratio": lower_wick / day_range if day_range else np.nan,
            }
        )
    return pd.concat([events.reset_index(drop=True), pd.DataFrame(rows)], axis=1)


def summarize_event_study(events: pd.DataFrame) -> pd.DataFrame:
    rows: list[dict[str, object]] = []
    for horizon in base.HORIZONS:
        baseline = float(events[f"T+{horizon}"].mean())
        lower25 = float(events[f"lower_wick_25_T+{horizon}"].mean())
        rows.append(
            {
                "horizon": f"T+{horizon}",
                "events": int(len(events)),
                "baseline_return": baseline,
                "lower_wick_25_return": lower25,
                "delta_vs_baseline": lower25 - baseline,
            }
        )
    return pd.DataFrame(rows)


def select_case_examples(events: pd.DataFrame) -> pd.DataFrame:
    filtered = events[events["lower_wick_25_filter"]].copy()
    avoided = filtered.sort_values(["T+5", "signal_date"], ascending=[True, False]).head(2).copy()
    avoided["case_type"] = "필터가 막는 손실"

    missed = filtered[filtered["T+5"].gt(0)].sort_values(
        ["T+5", "signal_date"],
        ascending=[False, False],
    ).head(2).copy()
    missed["case_type"] = "필터가 포기하는 수익"

    cases = pd.concat([avoided, missed], ignore_index=True)
    return cases.drop_duplicates(["signal_date", "symbol"]).head(4)


def make_case_figure(
    cases: pd.DataFrame,
    open_: pd.DataFrame,
    high: pd.DataFrame,
    low: pd.DataFrame,
    close: pd.DataFrame,
) -> Path:
    plt.rcParams["font.family"] = ["Malgun Gothic", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False
    fig, axes = plt.subplots(2, 2, figsize=(11.6, 7.0))
    flat_axes = list(axes.ravel())
    for ax, (_, event) in zip(flat_axes, cases.iterrows(), strict=False):
        plot_case(ax, event, open_, high, low, close)
    for ax in flat_axes[len(cases) :]:
        ax.axis("off")
    fig.suptitle("25% 필터 사례: 막아주는 손실과 포기하는 수익", x=0.02, y=1.02, ha="left", fontsize=14, color=base.NAVY, fontweight="bold")
    fig.text(
        0.02,
        0.975,
        "금색은 신호일 T. 25% 후보는 신호일 아래꼬리 조건 충족 시 T+1에 진입하지 않고 수익률 0으로 처리한다.",
        ha="left",
        color=base.SLATE,
        fontsize=9,
    )
    fig.tight_layout(rect=[0, 0, 1, 0.93], h_pad=2.0, w_pad=1.4)
    path = FIG_DIR / "lower_wick_25_cases.png"
    fig.savefig(path, dpi=190, bbox_inches="tight")
    plt.close(fig)
    return path


def plot_case(
    ax,
    event: pd.Series,
    open_: pd.DataFrame,
    high: pd.DataFrame,
    low: pd.DataFrame,
    close: pd.DataFrame,
) -> None:
    signal_date = pd.Timestamp(event["signal_date"])
    symbol = str(event["symbol"])
    index = open_.index
    pos = index.get_loc(signal_date)
    dates = index[max(0, pos - 4) : min(len(index), pos + 6)]
    x = np.arange(len(dates))
    o = open_.loc[dates, symbol].astype(float)
    h = high.loc[dates, symbol].astype(float)
    low_series = low.loc[dates, symbol].astype(float)
    c = close.loc[dates, symbol].astype(float)
    signal_x = int(np.where(dates == signal_date)[0][0])
    entry_date = pd.Timestamp(event["entry_date"])
    entry_x = int(np.where(dates == entry_date)[0][0]) if entry_date in set(dates) else signal_x + 1

    ax.axvspan(signal_x - 0.45, signal_x + 0.45, color=base.GOLD, alpha=0.18, zorder=0)
    ax.axvline(entry_x, color=base.CORAL, linestyle="--", linewidth=1.1, alpha=0.8)
    for i in x:
        color = base.CORAL if c.iloc[i] >= o.iloc[i] else base.TEAL
        ax.vlines(i, low_series.iloc[i], h.iloc[i], color=color, linewidth=1.15, alpha=0.95)
        bottom = min(o.iloc[i], c.iloc[i])
        height = max(abs(c.iloc[i] - o.iloc[i]), max(h.max() - low_series.min(), 1) * 0.006)
        ax.add_patch(Rectangle((i - 0.28, bottom), 0.56, height, facecolor=color, edgecolor=color, alpha=0.85))

    label_name = str(event.get("name") or "").strip()
    label = f"{label_name} {symbol}" if label_name else symbol
    lower_wick = float(event["lower_wick_ratio"]) * 100
    ax.set_title(f"{label}\n{signal_date:%Y-%m-%d} | 아래꼬리 {lower_wick:.1f}%", loc="left", fontsize=10.2, color=base.NAVY, fontweight="bold")
    case_type = str(event["case_type"])
    base_t5 = float(event["T+5"]) * 100
    note_color = base.TEAL if base_t5 > 0 else base.CORAL
    note = f"{case_type}\nbaseline T+5 {base_t5:+.1f}%"
    ax.text(0.03, 0.08, note, transform=ax.transAxes, fontsize=8.0, color=note_color, fontweight="bold")

    tick_labels = [d.strftime("%m/%d") if i in {0, signal_x, entry_x, len(dates) - 1} else "" for i, d in enumerate(dates)]
    ax.set_xticks(x, tick_labels, fontsize=7.6)
    ax.tick_params(axis="y", labelsize=7.6, colors=base.SLATE)
    ax.grid(axis="y", color=base.LINE, linewidth=0.7, alpha=0.6)
    for spine in ("top", "right", "left"):
        ax.spines[spine].set_visible(False)
    ax.spines["bottom"].set_color(base.LINE)


def write_workbook(data: LowerWick25ReportData) -> None:
    path = OUT_DIR / f"{REPORT_STEM}.xlsx"
    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        data.event_study.to_excel(writer, sheet_name="EVENT_STUDY_COMPARISON", index=False)
        data.cases.to_excel(writer, sheet_name="CASE_EXAMPLES", index=False)
        data.events.to_excel(writer, sheet_name="EVENTS_WITH_25_FILTER", index=False)


def build_docx(data: LowerWick25ReportData, path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    base.configure_styles(doc)
    set_docx_footer(section)

    base.add_title(doc, "아래꼬리 25% 미진입 후보")
    base.add_subtitle(doc, "신호일 T의 고정 조건만 사용한 간결 진단 보고서")

    base.add_heading(doc, "조건", level=1)
    base.add_table(
        doc,
        ["항목", "내용"],
        condition_rows(data.events),
        widths=[1.45, 4.85],
        font_size=8.5,
    )

    base.add_heading(doc, "경제적 함의", level=1)
    base.add_table(
        doc,
        ["함의", "설명"],
        lower_wick_rationale_rows(),
        widths=[1.45, 4.85],
        font_size=8.4,
    )

    base.add_heading(doc, "Event Study 비교", level=1)
    base.add_table(
        doc,
        event_study_headers(),
        event_study_rows(data.event_study),
        widths=[0.75, 1.0, 1.05, 1.25, 1.15, 1.05],
        font_size=7.9,
    )
    doc.add_picture(str(data.figures["event_study_matrix"]), width=Inches(6.05))
    base.last_para(doc).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(str(data.figures["event_study"]), width=Inches(5.95))
    base.last_para(doc).alignment = WD_ALIGN_PARAGRAPH.CENTER

    base.add_heading(doc, "사례", level=1)
    doc.add_picture(str(data.figures["cases"]), width=Inches(6.25))
    base.last_para(doc).alignment = WD_ALIGN_PARAGRAPH.CENTER

    base.add_heading(doc, "판단", level=1)
    base.add_note(
        doc,
        "25% 후보는 baseline 단독 적용 기준 T+1/T+2는 악화되지만 T+5는 +24.6bp 개선된다. 즉 단기 진입 필터라기보다 보유 기간이 길어질 때 손실 꼬리를 줄이는 후보로 보는 편이 맞다.",
    )
    doc.save(path)


def build_pdf(data: LowerWick25ReportData, path: Path) -> None:
    base.register_pdf_fonts()
    styles = base.pdf_styles()
    doc = SimpleDocTemplate(
        str(path),
        pagesize=landscape(letter),
        rightMargin=0.55 * inch,
        leftMargin=0.55 * inch,
        topMargin=0.48 * inch,
        bottomMargin=0.45 * inch,
        title=REPORT_STEM,
    )
    story: list[object] = [
        Paragraph("아래꼬리 25% 미진입 후보", styles["TitleKR"]),
        Paragraph("신호일 T의 고정 조건만 사용한 간결 진단 보고서", styles["SubtitleKR"]),
        Paragraph("조건", styles["H1KR"]),
        base.pdf_table([["항목", "내용"]] + condition_rows(data.events), [1.35 * inch, 8.45 * inch], styles),
        Paragraph("경제적 함의", styles["H1KR"]),
        base.pdf_table([["함의", "설명"]] + lower_wick_rationale_rows(), [1.55 * inch, 8.25 * inch], styles),
        Paragraph("Event Study 비교", styles["H1KR"]),
        base.pdf_table([event_study_headers()] + event_study_rows(data.event_study), [0.95 * inch, 1.45 * inch, 1.65 * inch, 1.45 * inch], styles, compact=True),
        Spacer(1, 0.05 * inch),
        base.pdf_image(data.figures["event_study_matrix"], width=6.6 * inch),
        Spacer(1, 0.05 * inch),
        base.pdf_image(data.figures["event_study"], width=4.65 * inch),
        PageBreak(),
        Paragraph("사례", styles["H1KR"]),
        base.pdf_image(data.figures["cases"], width=9.1 * inch),
        Spacer(1, 0.12 * inch),
        note_pdf(
            "판단: 25% 후보는 baseline 단독 적용 기준 T+1/T+2는 악화되지만 T+5는 +24.6bp 개선된다. 단기 성과 필터보다는 장기 보유 시 손실 꼬리를 줄이는 후보로 보는 편이 맞다.",
            styles,
        ),
    ]
    doc.build(story, onFirstPage=pdf_footer, onLaterPages=pdf_footer)


def condition_rows(events: pd.DataFrame) -> list[list[str]]:
    total_no_entry = int(events["combined_no_entry_25"].sum())
    active = int((~events["combined_no_entry_25"]).sum())
    return [
        ["대상", f"2016-01-01~2026-05-27 path-normal {len(events):,}건"],
        ["고정 조건", "신호일 T의 아래꼬리 비율 >= 25%이면 T+1 진입하지 않음"],
        ["계산식", "아래꼬리 비율 = (min(T 시가, T 종가) - T 저가) / (T 고가 - T 저가)"],
        ["미래참조", "없음. 신호일 T의 OHLC만 사용"],
        ["청산", "필터 미해당 이벤트는 baseline과 동일하게 T+1 시가 숏 진입 후 T+1~T+5 종가 청산"],
        ["건수", f"path-normal 전체 {len(events):,}건, 아래꼬리 25% 미진입 {total_no_entry:,}건, 실제 진입 {active:,}건"],
    ]


def event_study_headers() -> list[str]:
    return ["청산", "baseline", "아래꼬리 25%", "vs baseline"]


def event_study_rows(event_study: pd.DataFrame) -> list[list[str]]:
    rows: list[list[str]] = []
    for _, row in event_study.iterrows():
        rows.append(
            [
                str(row["horizon"]),
                fmt_pct(row["baseline_return"]),
                fmt_pct(row["lower_wick_25_return"]),
                fmt_bp(row["delta_vs_baseline"]),
            ]
        )
    return rows


def set_docx_footer(section) -> None:
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"KOSPI200 변동성 매도 전략 | {REPORT_STEM}")
    r.font.name = "Malgun Gothic"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(98, 110, 125)


def note_pdf(text: str, styles: dict[str, ParagraphStyle]):
    return base.note_pdf(text, styles)


def pdf_footer(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFont("Malgun", 7.5)
    canvas.setFillColor(colors.HexColor("#62707D"))
    canvas.drawRightString(10.55 * inch, 0.2 * inch, f"KOSPI200 변동성 매도 전략 | {REPORT_STEM} | {doc.page}")
    canvas.restoreState()


def fmt_pct(value: object) -> str:
    return f"{float(value) * 100:+.2f}%"


def fmt_bp(value: object) -> str:
    return f"{float(value) * 10000:+.1f}bp"


def lower_wick_rationale_rows() -> list[list[str]]:
    return [
        ["신호일 방어 매수", "긴 아래꼬리는 장중 급락을 매수세가 되받아낸 흔적이다. 급등 이후에도 저가 매수 수요가 남아 있으면 숏 진입의 기대값이 약해질 수 있다."],
        ["진입 전 필터", "baseline은 path-normal 356건 모두 T+1 시가에 진입하지만, 25% 필터는 신호일 저가 매수 방어가 강한 이벤트 91건을 진입 전 제거한다."],
        ["장기 보유 리스크 축소", "결과상 T+1/T+2는 둔화되지만 T+5 개선 폭이 커서, 짧은 청산보다 보유 기간이 길어질 때 손실 꼬리 관리 후보로 해석하는 편이 맞다."],
        ["기회비용 존재", "태광산업, 현대로템처럼 25% 조건에 걸리지만 실제로는 수익이 난 사례도 있어 단독 본전략 편입보다는 후보 필터로 두고 추가 검증이 필요하다."],
    ]


def make_event_study_figure(event_study: pd.DataFrame) -> Path:
    plt.rcParams["font.family"] = ["Malgun Gothic", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False
    fig, ax = plt.subplots(figsize=(8.7, 4.2))
    x = np.arange(len(event_study))
    baseline = event_study["baseline_return"] * 100
    lower25 = event_study["lower_wick_25_return"] * 100
    ax.bar(x, baseline, width=0.64, color=base.GRAY_BAR, edgecolor="none", label="baseline", zorder=2)
    ax.plot(x, lower25, marker="o", linewidth=2.7, color=base.CORAL, label="아래꼬리 25% 필터", zorder=5)
    ax.fill_between(x, lower25, baseline, color=base.CORAL, alpha=0.08, zorder=1)
    ax.axhline(0, color=base.LINE, linewidth=1.0)
    ax.set_xticks(x, event_study["horizon"])
    ax.set_ylabel("평균 수익률(%)", color=base.SLATE)
    ax.set_title("청산 시점별 평균 수익률", loc="left", color=base.NAVY, fontweight="bold")
    ax.legend(frameon=False, loc="upper right")
    base.style_axis(ax)
    y_low = min(event_study[["baseline_return", "lower_wick_25_return"]].min()) * 100
    y_high = max(event_study[["baseline_return", "lower_wick_25_return"]].max()) * 100
    pad = max((y_high - y_low) * 0.25, 0.16)
    ax.set_ylim(y_low - pad, y_high + pad)
    for i, value in enumerate(lower25):
        ax.text(i, value + 0.035, f"{value:+.2f}%", ha="center", va="bottom", fontsize=8.5, color=base.CORAL, fontweight="bold")
    fig.tight_layout()
    path = FIG_DIR / "event_study_lower_wick_25.png"
    fig.savefig(path, dpi=190, bbox_inches="tight")
    plt.close(fig)
    return path


def make_event_study_matrix_figure(event_study: pd.DataFrame) -> Path:
    plt.rcParams["font.family"] = ["Malgun Gothic", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False
    rows = []
    cell_colors = []
    for _, row in event_study.iterrows():
        rows.append(
            [
                row["horizon"],
                f"{row['baseline_return'] * 100:+.2f}%",
                f"{row['lower_wick_25_return'] * 100:+.2f}%",
                f"{row['delta_vs_baseline'] * 10000:+.1f}bp",
            ]
        )
        cell_colors.append(
            [
                "#FFFFFF",
                base.return_color(row["baseline_return"]),
                base.return_color(row["lower_wick_25_return"]),
                base.delta_color(row["delta_vs_baseline"]),
            ]
        )

    fig, ax = plt.subplots(figsize=(10.2, 3.55))
    ax.axis("off")
    ax.text(0.0, 1.12, "아래꼬리 25% 후보 Event Study 비교", transform=ax.transAxes, fontsize=14, fontweight="bold", color=base.NAVY)
    ax.text(
        0.0,
        1.03,
        "baseline은 path-normal 356건 전부 T+1 시가 숏 진입, 아래꼬리 25% 필터는 조건 충족 91건을 미진입 0으로 처리",
        transform=ax.transAxes,
        fontsize=9,
        color=base.SLATE,
    )
    table = ax.table(
        cellText=rows,
        colLabels=["청산", "baseline", "아래꼬리 25%", "vs baseline"],
        cellLoc="center",
        colLoc="center",
        cellColours=cell_colors,
        colColours=[base.NAVY] * 6,
        bbox=[0.0, 0.0, 1.0, 0.88],
    )
    base.style_matplotlib_table(table, header_color=base.NAVY, font_size=8.5)
    path = FIG_DIR / "event_study_lower_wick_25_matrix.png"
    fig.savefig(path, dpi=190, bbox_inches="tight")
    plt.close(fig)
    return path


if __name__ == "__main__":
    main()
