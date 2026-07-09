from __future__ import annotations

# ruff: noqa: E402

from dataclasses import dataclass
from pathlib import Path
import sys

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import landscape, letter
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from scripts import build_team_strat1_adjusted_report as base


OUT_DIR = base.OUT_DIR
FIG_DIR = OUT_DIR / "lower_wick_figures"
REPORT_STEM = "report_lower_wick"


@dataclass(frozen=True)
class LowerWickReportData:
    events: pd.DataFrame
    quartile_summary: pd.DataFrame
    filter_impact: pd.DataFrame
    threshold_sweep: pd.DataFrame
    figures: dict[str, Path]


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    data = build_report_data()
    write_workbook(data)
    docx_path = OUT_DIR / f"{REPORT_STEM}.docx"
    pdf_path = OUT_DIR / f"{REPORT_STEM}.pdf"
    build_docx(data, docx_path)
    build_pdf(data, pdf_path)
    print(f"wrote {docx_path}")
    print(f"wrote {pdf_path}")
    print(f"wrote {OUT_DIR / f'{REPORT_STEM}.xlsx'}")


def build_report_data() -> LowerWickReportData:
    adjusted = base.build_report_data()
    open_ = base.read_price("qw_adj_o")
    high = base.read_price("qw_adj_h").reindex_like(open_)
    low = base.read_price("qw_adj_l").reindex_like(open_)
    close = base.read_price("qw_adj_c").reindex_like(open_)

    events = adjusted.events.copy()
    events = add_lower_wick_metrics(events, open_, high, low, close)
    quartile_summary = summarize_quartiles(events)
    filter_impact = summarize_filter_impact(events)
    threshold_sweep = summarize_threshold_sweep(events)
    figures = make_figures(quartile_summary, filter_impact, threshold_sweep)
    return LowerWickReportData(events, quartile_summary, filter_impact, threshold_sweep, figures)


def add_lower_wick_metrics(
    events: pd.DataFrame,
    open_: pd.DataFrame,
    high: pd.DataFrame,
    low: pd.DataFrame,
    close: pd.DataFrame,
) -> pd.DataFrame:
    rows: list[dict[str, float]] = []
    for _, event in events.iterrows():
        date = pd.Timestamp(event["signal_date"])
        symbol = str(event["symbol"])
        open_price = float(open_.at[date, symbol])
        high_price = float(high.at[date, symbol])
        low_price = float(low.at[date, symbol])
        close_price = float(close.at[date, symbol])
        day_range = high_price - low_price
        body_low = min(open_price, close_price)
        body_high = max(open_price, close_price)
        lower_wick = max(body_low - low_price, 0.0)
        upper_wick = max(high_price - body_high, 0.0)
        body = abs(close_price - open_price)
        rows.append(
            {
                "day_range": day_range,
                "lower_wick": lower_wick,
                "upper_wick": upper_wick,
                "body": body,
                "lower_wick_ratio": lower_wick / day_range if day_range else np.nan,
                "upper_wick_ratio": upper_wick / day_range if day_range else np.nan,
                "body_ratio": body / day_range if day_range else np.nan,
                "close_position": (close_price - low_price) / day_range if day_range else np.nan,
            }
        )
    enriched = pd.concat([events.reset_index(drop=True), pd.DataFrame(rows)], axis=1)
    enriched["lower_wick_q"] = pd.qcut(
        enriched["lower_wick_ratio"],
        4,
        labels=["Q1 낮음", "Q2", "Q3", "Q4 높음"],
        duplicates="drop",
    )
    return enriched


def summarize_quartiles(events: pd.DataFrame) -> pd.DataFrame:
    rows: list[dict[str, object]] = []
    for quartile, subset in events.groupby("lower_wick_q", observed=True):
        row: dict[str, object] = {
            "quartile": str(quartile),
            "events": int(len(subset)),
            "lower_wick_avg": float(subset["lower_wick_ratio"].mean()),
            "not_entered_rate": float(subset["entry_status"].eq("not_entered").mean()),
            "stop_hit_T3": float(subset["prior_high_touch_T3"].mean()),
            "stop_hit_T5": float(subset["prior_high_touch_T5"].mean()),
        }
        for horizon in base.HORIZONS:
            baseline = subset[f"T+{horizon}"].mean()
            adjusted = subset[f"adjusted_T+{horizon}"].mean()
            row[f"baseline_T{horizon}"] = float(baseline)
            row[f"stop_T{horizon}"] = float(adjusted)
            row[f"delta_T{horizon}"] = float(adjusted - baseline)
        rows.append(row)
    return pd.DataFrame(rows)


def summarize_filter_impact(events: pd.DataFrame) -> pd.DataFrame:
    scenarios = [
        ("전체 358건", events),
        ("아래꼬리 Q4 제외", events[events["lower_wick_q"].ne("Q4 높음")]),
    ]
    rows: list[dict[str, object]] = []
    for label, subset in scenarios:
        row: dict[str, object] = {
            "scenario": label,
            "events": int(len(subset)),
            "excluded": int(len(events) - len(subset)),
            "not_entered_events": int(subset["entry_status"].eq("not_entered").sum()),
        }
        for horizon in base.HORIZONS:
            baseline = subset[f"T+{horizon}"].mean()
            adjusted = subset[f"adjusted_T+{horizon}"].mean()
            row[f"baseline_T{horizon}"] = float(baseline)
            row[f"stop_T{horizon}"] = float(adjusted)
            row[f"delta_T{horizon}"] = float(adjusted - baseline)
        rows.append(row)
    return pd.DataFrame(rows)


def summarize_threshold_sweep(events: pd.DataFrame) -> pd.DataFrame:
    rows: list[dict[str, object]] = []
    for threshold in (0.20, 0.25, 0.30, 0.35, 0.40):
        kept = events[events["lower_wick_ratio"].lt(threshold)]
        row: dict[str, object] = {
            "threshold": threshold,
            "rule": f"lower_wick_ratio < {threshold:.0%}",
            "kept_events": int(len(kept)),
            "excluded_events": int(len(events) - len(kept)),
        }
        for horizon in base.HORIZONS:
            row[f"stop_T{horizon}"] = float(kept[f"adjusted_T+{horizon}"].mean()) if len(kept) else np.nan
        rows.append(row)
    return pd.DataFrame(rows)


def make_figures(
    quartile_summary: pd.DataFrame,
    filter_impact: pd.DataFrame,
    threshold_sweep: pd.DataFrame,
) -> dict[str, Path]:
    plt.rcParams["font.family"] = ["Malgun Gothic", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False
    figures: dict[str, Path] = {}

    fig, ax = plt.subplots(figsize=(8.6, 4.3))
    x = np.arange(len(quartile_summary))
    ax.bar(x - 0.18, quartile_summary["stop_hit_T3"] * 100, width=0.36, color=base.CYAN, label="T+3 stop hit")
    ax.bar(x + 0.18, quartile_summary["stop_hit_T5"] * 100, width=0.36, color=base.TEAL, label="T+5 stop hit")
    ax.set_xticks(x, quartile_summary["quartile"])
    ax.set_ylabel("stop hit 비율(%)", color=base.SLATE)
    ax.set_title("아래꼬리 비율이 높을수록 stop hit가 증가", loc="left", color=base.NAVY, fontweight="bold")
    ax.legend(frameon=False, loc="upper left")
    base.style_axis(ax)
    fig.tight_layout()
    path = FIG_DIR / "lower_wick_stop_hit_by_quartile.png"
    fig.savefig(path, dpi=190, bbox_inches="tight")
    plt.close(fig)
    figures["stop_hit_by_quartile"] = path

    fig, ax = plt.subplots(figsize=(8.6, 4.3))
    x = np.arange(len(base.HORIZONS))
    for _, row in filter_impact.iterrows():
        ax.plot(x, [row[f"stop_T{h}"] * 100 for h in base.HORIZONS], marker="o", linewidth=2.4, label=str(row["scenario"]))
    ax.axhline(0, color=base.LINE, linewidth=1.0)
    ax.set_xticks(x, [f"T+{h}" for h in base.HORIZONS])
    ax.set_ylabel("stop 적용 후 평균 수익률(%)", color=base.SLATE)
    ax.set_title("Q4 제외 시 stop 적용 후 수익률 변화", loc="left", color=base.NAVY, fontweight="bold")
    ax.legend(frameon=False, loc="upper left")
    base.style_axis(ax)
    fig.tight_layout()
    path = FIG_DIR / "lower_wick_filter_impact.png"
    fig.savefig(path, dpi=190, bbox_inches="tight")
    plt.close(fig)
    figures["filter_impact"] = path

    fig, ax = plt.subplots(figsize=(8.6, 4.0))
    labels = [f"< {threshold:.0%}" for threshold in threshold_sweep["threshold"]]
    ax.bar(labels, threshold_sweep["stop_T5"] * 100, color=base.TEAL, alpha=0.86)
    ax.axhline(0, color=base.LINE, linewidth=1.0)
    ax.set_ylabel("T+5 stop 평균 수익률(%)", color=base.SLATE)
    ax.set_title("아래꼬리 임계값별 T+5 수익률", loc="left", color=base.NAVY, fontweight="bold")
    base.style_axis(ax)
    fig.tight_layout()
    path = FIG_DIR / "lower_wick_threshold_sweep.png"
    fig.savefig(path, dpi=190, bbox_inches="tight")
    plt.close(fig)
    figures["threshold_sweep"] = path

    return figures


def write_workbook(data: LowerWickReportData) -> None:
    path = OUT_DIR / f"{REPORT_STEM}.xlsx"
    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        data.quartile_summary.to_excel(writer, sheet_name="QUARTILE_SUMMARY", index=False)
        data.filter_impact.to_excel(writer, sheet_name="FILTER_IMPACT", index=False)
        data.threshold_sweep.to_excel(writer, sheet_name="THRESHOLD_SWEEP", index=False)
        data.events.to_excel(writer, sheet_name="EVENTS_WITH_LOWER_WICK", index=False)


def build_docx(data: LowerWickReportData, path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    base.configure_styles(doc)
    base.set_footer(section)

    base.add_title(doc, "KOSPI200 유성형 매도 전략")
    base.add_subtitle(doc, "아래꼬리 필터 후보 진단 보고서")
    base.add_heading(doc, "검토 목적", level=1)
    base.add_table(
        doc,
        ["항목", "내용"],
        [
            ["대상", "2016~현재 dedup 358건 유성형 매도 이벤트"],
            ["가설", "신호일 아래꼬리가 길면 저가 매수 방어가 강해 숏 진입 후 되돌림 위험이 커질 수 있다."],
            ["지표", "아래꼬리 비율 = (min(시가, 종가) - 저가) / (고가 - 저가)"],
            ["검토 방식", "아래꼬리 비율 4분위별 baseline/stop 수익률, stop hit, Q4 제외 효과를 비교"],
        ],
        widths=[1.45, 4.85],
    )

    base.add_heading(doc, "핵심 결과", level=1)
    base.add_table(
        doc,
        ["결과", "해석"],
        [
            ["stop hit 증가", "아래꼬리 Q4의 T+5 stop hit는 55.6%로 Q1 28.9% 대비 높다."],
            ["Q4 제외 효과", "Q4를 제외하면 T+5 stop 수익률이 +0.17%에서 +0.39%로 개선된다."],
            ["단독 필터 한계", "Q4 내부에도 태광산업, 신풍제약, LIG넥스원 같은 수익 사례가 있어 전면 제외는 손실 가능성이 있다."],
            ["권고", "즉시 하드 필터보다 아래꼬리 + 종가 위치 + T+1 시가 강도를 결합한 조건부 필터로 발전시키는 편이 낫다."],
        ],
        widths=[1.45, 4.85],
    )

    base.add_heading(doc, "아래꼬리 분위수별 성과", level=1)
    base.add_table(doc, quartile_docx_rows(data.quartile_summary)[0], quartile_docx_rows(data.quartile_summary)[1], widths=[1.0, 0.7, 1.0, 1.0, 1.0, 1.0, 1.0], font_size=7.6)
    doc.add_picture(str(data.figures["stop_hit_by_quartile"]), width=Inches(5.9))
    base.last_para(doc).alignment = base.WD_ALIGN_PARAGRAPH.CENTER

    base.add_heading(doc, "Q4 제외 효과", level=1)
    base.add_table(doc, filter_docx_rows(data.filter_impact)[0], filter_docx_rows(data.filter_impact)[1], widths=[1.25, 0.7, 0.85, 0.85, 0.85, 0.85, 0.85], font_size=7.6)
    doc.add_picture(str(data.figures["filter_impact"]), width=Inches(5.9))
    base.last_para(doc).alignment = base.WD_ALIGN_PARAGRAPH.CENTER

    base.add_heading(doc, "임계값 민감도", level=1)
    base.add_table(doc, threshold_docx_rows(data.threshold_sweep)[0], threshold_docx_rows(data.threshold_sweep)[1], widths=[1.25, 0.8, 0.8, 0.9, 0.9, 0.9, 0.9], font_size=7.6)
    doc.add_picture(str(data.figures["threshold_sweep"]), width=Inches(5.7))
    base.last_para(doc).alignment = base.WD_ALIGN_PARAGRAPH.CENTER

    base.add_heading(doc, "판단", level=1)
    base.add_note(
        doc,
        "아래꼬리 Q4 제외는 보유기간이 길수록 성과 개선 여지가 있지만, 강한 수익 사례도 함께 제거한다. 따라서 현재 보고서 본전략에 바로 편입하기보다는 조건부 필터 후보로 두고 종가 위치, T+1 시가 강도, 거래량 유지 여부를 함께 검증하는 것이 적절하다.",
    )
    doc.save(path)


def build_pdf(data: LowerWickReportData, path: Path) -> None:
    base.register_pdf_fonts()
    styles = base.pdf_styles()
    doc = SimpleDocTemplate(
        str(path),
        pagesize=landscape(letter),
        rightMargin=0.55 * inch,
        leftMargin=0.55 * inch,
        topMargin=0.48 * inch,
        bottomMargin=0.45 * inch,
    )
    story = [
        Paragraph("KOSPI200 유성형 매도 전략", styles["TitleKR"]),
        Paragraph("아래꼬리 필터 후보 진단 보고서", styles["SubtitleKR"]),
        Paragraph("검토 목적", styles["H1KR"]),
        base.pdf_table(
            [["항목", "내용"]]
            + [
                ["대상", "2016~현재 dedup 358건 유성형 매도 이벤트"],
                ["가설", "신호일 아래꼬리가 길면 저가 매수 방어가 강해 숏 진입 후 되돌림 위험이 커질 수 있다."],
                ["지표", "아래꼬리 비율 = (min(시가, 종가) - 저가) / (고가 - 저가)"],
                ["검토 방식", "아래꼬리 비율 4분위별 baseline/stop 수익률, stop hit, Q4 제외 효과를 비교"],
            ],
            [1.35 * inch, 8.45 * inch],
            styles,
        ),
        Paragraph("핵심 결과", styles["H1KR"]),
        base.pdf_table(
            [["결과", "해석"]]
            + [
                ["stop hit 증가", "아래꼬리 Q4의 T+5 stop hit는 55.6%로 Q1 28.9% 대비 높다."],
                ["Q4 제외 효과", "Q4를 제외하면 T+5 stop 수익률이 +0.17%에서 +0.39%로 개선된다."],
                ["단독 필터 한계", "Q4 내부에도 수익 사례가 있어 전면 제외는 좋은 이벤트를 같이 제거할 수 있다."],
                ["권고", "하드 필터보다 아래꼬리 + 종가 위치 + T+1 시가 강도를 결합한 조건부 필터로 발전시키는 편이 낫다."],
            ],
            [1.6 * inch, 8.2 * inch],
            styles,
        ),
        PageBreak(),
        Paragraph("아래꼬리 분위수별 성과", styles["H1KR"]),
        base.pdf_table([quartile_docx_rows(data.quartile_summary)[0]] + quartile_docx_rows(data.quartile_summary)[1], [1.15 * inch, 0.7 * inch, 1.05 * inch, 1.1 * inch, 1.05 * inch, 1.05 * inch, 1.05 * inch], styles, compact=True),
        Spacer(1, 0.12 * inch),
        base.pdf_image(data.figures["stop_hit_by_quartile"], width=7.1 * inch),
        PageBreak(),
        Paragraph("Q4 제외 효과", styles["H1KR"]),
        base.pdf_table([filter_docx_rows(data.filter_impact)[0]] + filter_docx_rows(data.filter_impact)[1], [1.4 * inch, 0.75 * inch, 1.0 * inch, 1.0 * inch, 1.0 * inch, 1.0 * inch, 1.0 * inch], styles, compact=True),
        Spacer(1, 0.12 * inch),
        base.pdf_image(data.figures["filter_impact"], width=7.2 * inch),
        PageBreak(),
        Paragraph("임계값 민감도", styles["H1KR"]),
        base.pdf_table([threshold_docx_rows(data.threshold_sweep)[0]] + threshold_docx_rows(data.threshold_sweep)[1], [1.3 * inch, 0.85 * inch, 0.85 * inch, 1.0 * inch, 1.0 * inch, 1.0 * inch, 1.0 * inch], styles, compact=True),
        Spacer(1, 0.12 * inch),
        base.pdf_image(data.figures["threshold_sweep"], width=7.0 * inch),
        Spacer(1, 0.12 * inch),
        base.note_pdf(
            "판단: 아래꼬리 Q4 제외는 보유기간이 길수록 성과 개선 여지가 있지만, 강한 수익 사례도 함께 제거한다. 현재 본전략에 바로 편입하기보다 조건부 필터 후보로 두고 종가 위치, T+1 시가 강도, 거래량 유지 여부를 함께 검증하는 것이 적절하다.",
            styles,
        ),
    ]
    doc.build(story, onFirstPage=base.pdf_footer, onLaterPages=base.pdf_footer)


def quartile_docx_rows(summary: pd.DataFrame) -> tuple[list[str], list[list[str]]]:
    header = ["구간", "n", "아래꼬리", "T+3 stop", "T+3 개선", "T+5 stop", "T+5 hit"]
    rows = []
    for _, row in summary.iterrows():
        rows.append(
            [
                row["quartile"],
                f"{int(row['events']):,}",
                base.pct(row["lower_wick_avg"]),
                base.pct(row["stop_T3"]),
                base.pct(row["delta_T3"], signed=True),
                base.pct(row["stop_T5"]),
                base.pct(row["stop_hit_T5"]),
            ]
        )
    return header, rows


def filter_docx_rows(summary: pd.DataFrame) -> tuple[list[str], list[list[str]]]:
    header = ["시나리오", "n", "T+1 stop", "T+2 stop", "T+3 stop", "T+4 stop", "T+5 stop"]
    rows = []
    for _, row in summary.iterrows():
        rows.append(
            [
                row["scenario"],
                f"{int(row['events']):,}",
                base.pct(row["stop_T1"]),
                base.pct(row["stop_T2"]),
                base.pct(row["stop_T3"]),
                base.pct(row["stop_T4"]),
                base.pct(row["stop_T5"]),
            ]
        )
    return header, rows


def threshold_docx_rows(summary: pd.DataFrame) -> tuple[list[str], list[list[str]]]:
    header = ["규칙", "유지", "제외", "T+2 stop", "T+3 stop", "T+4 stop", "T+5 stop"]
    rows = []
    for _, row in summary.iterrows():
        rows.append(
            [
                row["rule"],
                f"{int(row['kept_events']):,}",
                f"{int(row['excluded_events']):,}",
                base.pct(row["stop_T2"]),
                base.pct(row["stop_T3"]),
                base.pct(row["stop_T4"]),
                base.pct(row["stop_T5"]),
            ]
        )
    return header, rows


if __name__ == "__main__":
    main()
